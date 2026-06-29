



import os
import fitz  # PyMuPDF
from flask import Blueprint, request, jsonify
from app.models.knowledge_base import KnowledgeBase
from app.database.DatabaseOperationPostgreSQL import db_session
from flask_jwt_extended import jwt_required, get_jwt_identity, get_jwt
from qdrant_client import QdrantClient
from qdrant_client.models import PointStruct, VectorParams, Distance
from openai import OpenAI
import time
import requests
import random
from typing import List, Optional, Tuple
import json
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlsplit, urlunsplit
from app.models import BaseAgent, CustomBot
# Add docx import at the top level
import re
from logging_config import setup_logging
import pandas as pd
import numpy as np
import uuid
from sentence_transformers import SentenceTransformer
from sklearn.preprocessing import normalize
from app.models.suppliers_details import SupplierDetails  # Assuming your ORM model matches tbl_suppliers_detailsfrom .utils import (chunk_text, process_pdf, 
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. DOCX files will not be supported.")

# Initialize Qdrant client
QDRANT_URL = os.getenv("QDRANT_URL", "http://localhost:6333")
qdrant = QdrantClient(url=QDRANT_URL, timeout=120)

# Default embedding model
DEFAULT_EMBEDDING_MODEL = 'text-embedding-3-large'
DEFAULT_PROVIDER = 'openai'

knowledge_base_blueprint = Blueprint("knowledge_base", __name__)



logger = setup_logging("knowledge-base-routes", level="DEBUG")


def _normalize_url_for_crawl(url: str) -> str:
    """Normalize a URL for crawling and duplicate detection."""
    parsed = urlsplit(url.strip())
    scheme = parsed.scheme.lower()
    netloc = parsed.netloc.lower()
    path = parsed.path or "/"
    if path != "/":
        path = path.rstrip("/") or "/"
    return urlunsplit((scheme, netloc, path, parsed.query, parsed.fragment))


def _domain_variants(url: str) -> List[str]:
    """Return stable same-site host variants for safe cross-host crawl continuity."""
    host = (urlsplit(url).netloc or "").lower().strip()
    if not host:
        return []
    variants = {host}
    if host.startswith("www."):
        variants.add(host[4:])
    else:
        variants.add(f"www.{host}")
    return sorted(variants)


def _is_placeholder_crawl_url(url: str) -> bool:
    """Detect obvious placeholder or local test URLs that usually produce no crawlable content."""
    host = (urlsplit(url).netloc or "").lower().strip()
    if host.startswith("www."):
        host = host[4:]
    return host in {
        "example.com",
        "example.org",
        "example.net",
        "localhost",
        "127.0.0.1",
        "0.0.0.0",
    }


def _resolve_redirect_url(url: str, timeout: int = 15) -> Tuple[str, bool]:
    """Resolve HTTP redirects so crawling starts from the final destination."""
    headers = {"User-Agent": "Mozilla/5.0"}
    normalized_url = _normalize_url_for_crawl(url)

    try:
        response = requests.head(
            normalized_url,
            allow_redirects=True,
            timeout=timeout,
            headers=headers,
        )
        final_url = response.url or normalized_url
        return _normalize_url_for_crawl(final_url), final_url != normalized_url
    except Exception:
        try:
            response = requests.get(
                normalized_url,
                allow_redirects=True,
                timeout=timeout,
                headers=headers,
                stream=True,
            )
            final_url = response.url or normalized_url
            response.close()
            return _normalize_url_for_crawl(final_url), final_url != normalized_url
        except Exception as e:
            logger.warning(f"Unable to resolve redirects for {url}: {e}")
            return normalized_url, False


def _is_useful_scraped_content(content: str, url: str, min_chars: int = 200) -> bool:
    """Reject obviously low-value scrape output before chunking."""
    if not content or not content.strip():
        logger.warning(f"Skipping empty scraped content from {url}")
        return False

    cleaned = content.strip()
    if len(cleaned) < min_chars:
        logger.warning(f"Skipping short scraped content from {url}: {len(cleaned)} chars")
        return False

    words = re.findall(r"\w+", cleaned.lower())
    if len(words) < 40:
        logger.warning(f"Skipping low-word-count scraped content from {url}: {len(words)} words")
        return False

    unique_ratio = len(set(words)) / max(len(words), 1)
    if unique_ratio < 0.2:
        logger.warning(f"Skipping repetitive scraped content from {url}")
        return False

    return True


def _extract_domain_tokens(host: str) -> set:
    host = (host or "").lower()
    host = host[4:] if host.startswith("www.") else host
    parts = [p for p in host.split(".") if p and p not in {"com", "org", "net", "io", "co", "in", "us"}]
    return set(parts)


def _pre_analyze_crawl_target(seed_url: str, headers: dict, timeout: int = 12) -> dict:
    """
    Lightweight pre-analysis to select safe crawl mode/flags.
    Keeps requests limited to avoid delaying the real crawl.
    """
    normalized_seed = _normalize_url_for_crawl(seed_url)
    parsed_seed = urlsplit(normalized_seed)
    base_host = parsed_seed.netloc.lower()
    base_tokens = _extract_domain_tokens(base_host)
    result = {
        "seed_url": normalized_seed,
        "final_url": normalized_seed,
        "base_host": base_host,
        "final_host": base_host,
        "status_code": None,
        "redirect_chain": [normalized_seed],
        "related_domains": sorted(set(_domain_variants(normalized_seed))),
        "internal_link_count": 0,
        "external_related_link_count": 0,
        "link_density": 0,
        "query_link_count": 0,
        "duplicate_link_signals": 0,
        "is_js_heavy": False,
        "script_tag_count": 0,
        "text_len": 0,
        "selected_mode": "STANDARD",
        "flags": ["DEDUP_STRICT", "REDIRECT_CHAIN_TRACK"],
        "internal_links_sample": [],
    }

    page_html = ""
    final_url = normalized_seed
    try:
        resp = requests.get(
            normalized_seed,
            allow_redirects=True,
            timeout=timeout,
            headers={**headers, "User-Agent": "Mozilla/5.0"},
        )
        result["status_code"] = resp.status_code
        final_url = _normalize_url_for_crawl(resp.url or normalized_seed)
        result["final_url"] = final_url
        result["final_host"] = urlsplit(final_url).netloc.lower()
        chain = [h.headers.get("Location") for h in (resp.history or []) if h.headers.get("Location")]
        if chain:
            normalized_chain = [normalized_seed]
            for loc in chain:
                try:
                    normalized_chain.append(_normalize_url_for_crawl(requests.compat.urljoin(normalized_seed, loc)))
                except Exception:
                    continue
            normalized_chain.append(final_url)
            result["redirect_chain"] = normalized_chain
        elif final_url != normalized_seed:
            result["redirect_chain"] = [normalized_seed, final_url]
        page_html = resp.text or ""
    except Exception as e:
        logger.warning(f"[KB PRE-ANALYSIS] request failed for {normalized_seed}: {e}")

    # redirect signals in HTML
    html_lower = page_html.lower()
    if "http-equiv=\"refresh\"" in html_lower or "window.location" in html_lower or "location.href" in html_lower:
        if "REDIRECT_CHAIN_TRACK" not in result["flags"]:
            result["flags"].append("REDIRECT_CHAIN_TRACK")

    hrefs = re.findall(r'href=["\']([^"\']+)["\']', page_html, flags=re.IGNORECASE)
    resolved_links = []
    for href in hrefs:
        href = (href or "").strip()
        if not href or href.startswith("#") or href.startswith("mailto:") or href.startswith("tel:") or href.startswith("javascript:"):
            continue
        absolute = requests.compat.urljoin(final_url, href)
        if absolute.startswith("http://") or absolute.startswith("https://"):
            resolved_links.append(_normalize_url_for_crawl(absolute))

    unique_links = sorted(set(resolved_links))
    result["internal_links_sample"] = unique_links[:250]
    result["link_density"] = len(unique_links)
    result["query_link_count"] = len([u for u in unique_links if "?" in u])
    result["duplicate_link_signals"] = max(0, len(resolved_links) - len(unique_links))

    seed_hosts = set(_domain_variants(normalized_seed)) | set(_domain_variants(final_url))
    related_domains = set(seed_hosts)
    for link in unique_links:
        host = urlsplit(link).netloc.lower()
        if not host:
            continue
        if host in seed_hosts:
            result["internal_link_count"] += 1
            related_domains.add(host)
            continue
        host_tokens = _extract_domain_tokens(host)
        if base_tokens and host_tokens and len(base_tokens.intersection(host_tokens)) > 0:
            result["external_related_link_count"] += 1
            related_domains.add(host)

    result["related_domains"] = sorted(related_domains)
    result["script_tag_count"] = len(re.findall(r"<script\\b", page_html, flags=re.IGNORECASE))
    text_for_ratio = re.sub(r"<[^>]+>", " ", page_html)
    text_for_ratio = re.sub(r"\\s+", " ", text_for_ratio).strip()
    result["text_len"] = len(text_for_ratio)
    result["is_js_heavy"] = result["script_tag_count"] >= 20 and result["text_len"] < 4000

    # Mode selection
    if len(result["redirect_chain"]) > 1:
        result["selected_mode"] = "REDIRECT_RESOLVED"
    if len(result["related_domains"]) > 2 or result["external_related_link_count"] > 3:
        result["selected_mode"] = "MULTI_DOMAIN_GRAPH"
    if result["link_density"] >= 150:
        result["selected_mode"] = "DEEP_LINK"
    if result["is_js_heavy"]:
        result["selected_mode"] = "JS_RENDERED"

    if result["query_link_count"] > 12 or result["selected_mode"] in {"DEEP_LINK", "MULTI_DOMAIN_GRAPH"}:
        if "PAGINATION_SAFE" not in result["flags"]:
            result["flags"].append("PAGINATION_SAFE")

    return result

def get_embedding_model(provider: Optional[str] = None, model_name: Optional[str] = None, secret_key: Optional[str] = None) -> dict:
    """Return the embedding model configuration and vector size based on user input or default."""
    OPENAI_MODEL_DIMENSIONS = {
        "text-embedding-ada-002": 1536,
        "text-embedding-3-small": 1536,
        "text-embedding-3-large": 3072
    }

    if provider and model_name and secret_key:
        try:
            if provider.lower() == "openai":
                client = OpenAI(api_key=secret_key)
                vector_size = OPENAI_MODEL_DIMENSIONS.get(
                    model_name,
                    OPENAI_MODEL_DIMENSIONS[DEFAULT_EMBEDDING_MODEL],
                )
                logger.info(f"Using user-provided OpenAI embedding model: {model_name} (vector_size={vector_size})")
                return {
                    "provider": "openai",
                    "model_name": model_name,
                    "client": client,
                    "vector_size": vector_size
                }
            else:
                logger.warning(f"Unsupported provider '{provider}'. Falling back to default.")
        except Exception as e:
            logger.error(f"Failed to initialize user-provided model: {str(e)}")
            raise ValueError(f"Invalid provider or secret key: {str(e)}")

    try:
        openai_api_key = os.getenv("OPENAI_API_KEY")
        if not openai_api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set.")
        client = OpenAI(api_key=openai_api_key)
        vector_size = OPENAI_MODEL_DIMENSIONS.get(
            DEFAULT_EMBEDDING_MODEL,
            OPENAI_MODEL_DIMENSIONS["text-embedding-3-large"],
        )
        logger.info(f"Using default OpenAI embedding model: {DEFAULT_EMBEDDING_MODEL} (vector_size={vector_size})")
        return {
            "provider": "openai",
            "model_name": DEFAULT_EMBEDDING_MODEL,
            "client": client,
            "vector_size": vector_size
        }
    except Exception as e:
        logger.error(f"Failed to initialize default OpenAI model: {str(e)}")
        raise ValueError(f"Cannot initialize embedding model: {str(e)}")

def get_embeddings(texts: List[str], embedding_config: dict) -> List[List[float]]:
    """Generate embeddings for a list of texts using the specified embedding configuration."""
    try:
        if embedding_config["provider"] == "openai":
            response = embedding_config["client"].embeddings.create(
                model=embedding_config["model_name"],
                input=texts
            )
            return [item.embedding for item in response.data]
        else:
            raise ValueError(f"Unsupported provider: {embedding_config['provider']}")
    except Exception as e:
        logger.error(f"Error generating embeddings: {str(e)}")
        raise RuntimeError(f"Failed to generate embeddings: {str(e)}")

def create_qdrant_collection(collection_name: str, vector_size: int):
    """Create or recreate a Qdrant collection if it doesn't exist or has an incorrect vector size."""
    logger.info(f"Creating collection '{collection_name}' in Qdrant with vector_size={vector_size}...")
    try:
        collections_response = qdrant.get_collections()
        for collection in collections_response.collections:
            if collection.name == collection_name:
                collection_info = qdrant.get_collection(collection_name)
                current_vector_size = collection_info.config.params.vectors.size
                if current_vector_size != vector_size:
                    logger.warning(f"Collection '{collection_name}' has vector size {current_vector_size}, expected {vector_size}. Recreating.")
                    qdrant.delete_collection(collection_name)
                    break
                else:
                    logger.info(f"Collection '{collection_name}' already exists with correct vector size {vector_size}.")
                    return

        qdrant.create_collection(
            collection_name=collection_name,
            vectors_config=VectorParams(
                size=vector_size,
                distance=Distance.COSINE
            )
        )
        logger.info(f"Collection '{collection_name}' created successfully.")
    except Exception as e:
        logger.error(f"Error creating collection: {str(e)}")
        raise RuntimeError(f"Failed to create Qdrant collection: {str(e)}")

def create_directory_structure(knowledge_base_name: str, file_name: str) -> str:
    """Create directory structure for storing processed files."""
    base_path = os.path.join("knowledge_bases", knowledge_base_name, file_name)
    os.makedirs(base_path, exist_ok=True)
    return base_path

def create_website_chunks_directory(knowledge_base_name: str, website_name: str) -> str:
    """Create directory for storing website chunks."""
    website_dir = website_name.replace("https://", "").replace("http://", "").replace("/", "_")
    base_path = os.path.join("knowledge_bases", knowledge_base_name, "website_chunks", website_dir)
    os.makedirs(base_path, exist_ok=True)
    return base_path

def create_raw_text_chunks_directory(knowledge_base_name: str) -> str:
    """Create directory for storing raw text chunks."""
    timestamp = int(time.time())
    base_path = os.path.join("knowledge_bases", knowledge_base_name, "raw_text_chunks", f"text_{timestamp}")
    os.makedirs(base_path, exist_ok=True)
    return base_path

def insert_chunks_to_qdrant(collection_name: str, chunks: List[dict], embedding_config: dict, batch_size: int = 64, max_retries: int = 5):
    """Insert chunks into the specified Qdrant collection."""
    points = []
    for i in range(0, len(chunks), batch_size):
        batch_chunks = chunks[i:i + batch_size]
        try:
            texts = [chunk["text"] for chunk in batch_chunks]
            vectors = get_embeddings(texts, embedding_config)
            if not vectors or len(vectors) != len(texts):
                logger.error(f"Invalid embeddings generated for batch {i//batch_size + 1}")
                continue

            for chunk, vector in zip(batch_chunks, vectors):
                point = PointStruct(
                    id=str(uuid.uuid4()),  # ✅ ALWAYS UNIQUE
                    vector=vector,
                    payload=chunk
                )
                points.append(point)
        except Exception as e:
            logger.error(f"Error generating embeddings for batch {i//batch_size + 1}: {str(e)}")
            continue

    if not points:
        logger.error("No valid embeddings generated for any chunks")
        raise RuntimeError("Failed to generate any valid embeddings")

    for attempt in range(max_retries):
        try:
            for i in range(0, len(points), batch_size):
                batch = points[i:i + batch_size]
                qdrant.upsert(collection_name=collection_name, points=batch, wait=True)
                logger.info(f"Inserted batch {i//batch_size + 1} of {(len(points) + batch_size - 1)//batch_size} into Qdrant.")
            logger.info(f"Successfully inserted {len(chunks)} chunks into Qdrant collection '{collection_name}'.")
            return
        except Exception as e:
            if "Vector dimension error" in str(e):
                logger.error(f"Dimension mismatch detected: {str(e)}")
                raise ValueError(f"Qdrant collection expects different vector size: {str(e)}")
            logger.error(f"Attempt {attempt + 1}: Error inserting chunks - {str(e)}")
            time.sleep(2 ** attempt + random.uniform(0, 1))
    logger.error("Failed to insert chunks after multiple retries")
    raise RuntimeError("Failed to insert chunks into Qdrant after retries")

def get_text_chunks(file_path: str, chunk_chars: int, overlap: int) -> List[str]:
    """Extract text from a PDF and split into chunks."""
    try:
        doc = fitz.open(file_path)
        text = "".join([page.get_text() for page in doc])
        doc.close()
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_chars
            chunks.append(text[start:end])
            start = end - overlap
        return chunks
    except Exception as e:
        logger.error(f"Error extracting text chunks from {file_path}: {str(e)}")
        return []

def process_docx_file(file_path: str) -> str:
    """Extract text from a DOCX file."""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")
    
    try:
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text.strip())
        
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
        raise

def process_txt_file(file_path: str) -> str:
    """Extract text from a TXT file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {str(e)}")
            raise
    except Exception as e:
        logger.error(f"Error processing TXT file {file_path}: {str(e)}")
        raise

def _preview_text_for_log(text: str, max_chars: int = 1000) -> str:
    """Normalize text for compact log previews."""
    preview = re.sub(r"\s+", " ", text or "").strip()
    if len(preview) <= max_chars:
        return preview
    return preview[:max_chars] + " ... [truncated]"

def process_pdf(file_path: str, output_dir: str, knowledge_base_name: str, chunk_size: int, overlap: int) -> tuple:
    """Process a PDF file, extract text and images, and generate chunks."""
    try:
        pdf_document = fitz.open(file_path)
        extracted_pages = []

        for page_num, page in enumerate(pdf_document, start=1):
            page_text = page.get_text().strip()
            if page_text:
                extracted_pages.append(page_text)
                logger.info(
                    f"[PDF TEXT] {os.path.basename(file_path)} | page={page_num} | chars={len(page_text)} | preview={_preview_text_for_log(page_text, 500)}"
                )
            else:
                logger.info(f"[PDF TEXT] {os.path.basename(file_path)} | page={page_num} | chars=0 | preview=<empty>")

            image_list = page.get_images(full=True)
            for img_index, img in enumerate(image_list, start=1):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                image_filename = os.path.join(output_dir, f"page_{page_num}_image_{img_index}.{image_ext}")
                with open(image_filename, "wb") as image_file:
                    image_file.write(image_bytes)

        full_text = "\n\n".join(extracted_pages).strip()
        logger.info(
            f"[PDF TEXT] {os.path.basename(file_path)} | extracted_chars={len(full_text)} | pages_with_text={len(extracted_pages)}"
        )
        if full_text:
            logger.debug(f"[PDF TEXT RAW] {os.path.basename(file_path)} | preview={_preview_text_for_log(full_text, 4000)}")

        chunks = []
        start = 0
        while start < len(full_text):
            end = start + chunk_size
            chunks.append(full_text[start:end])
            start = end - overlap

        pdf_document.close()
        return True, f"Processed and stored PDF content for '{knowledge_base_name}'", chunks
    except Exception as e:
        logger.error(f"Error processing PDF {file_path}: {str(e)}")
        return False, str(e), []

def extract_attributes_from_text(text) -> dict:
    """
    Generic, safe attribute extractor.
    Accepts ANY input type.
    Never throws.
    """
    attributes = {}

    if not isinstance(text, str):
        return attributes

    text = text.strip()
    if not text:
        return attributes

    # Key: Value detection
    kv_pairs = re.findall(r"([\w\s]+)\s*:\s*([^\n]+)", text)
    if kv_pairs:
        attributes["key_values"] = {
            k.strip().lower().replace(" ", "_"): v.strip()  # ✅ FIX: Replace spaces with underscores
            for k, v in kv_pairs
        }

    # Identifier-like tokens
    tokens = re.findall(r"\b[A-Z0-9][A-Z0-9\-]{4,}\b", text)
    if tokens:
        attributes["identifiers"] = list(set(tokens))[:10]

    return attributes

def process_excel_file(file_path, chunk_size, chunk_overlap, batch_size=500):
    """
    Production-grade Excel → Text → Chunk extraction (NO embeddings here).

    Responsibilities:
        ✔ Dynamic column discovery
        ✔ Data cleaning + normalization
        ✔ Chunk creation with metadata
        ✔ Error safe, logging optimized
        ❌ Embedding generation (handled later by pipeline)
    
    Returns:
        (success: bool, message: str, chunks: List[dict])
    """

    try:
        # -------- Step 1: Load Safely --------
        try:
            df = pd.read_excel(file_path, dtype=str)
        except ValueError as e:
            return False, f"Unable to read Excel (format issue): {str(e)}", []
        except Exception as e:
            logger.exception("[EXCEL] Critical parsing failure")
            return False, f"Error parsing Excel: {str(e)}", []

        if df.empty:
            return False, "Excel contains no usable rows", []

        logger.info(f"[EXCEL] Loaded {file_path} → Rows: {len(df)}, Columns: {df.columns.tolist()}")

        # -------- Step 2: Detect Text Columns --------
        text_columns = df.select_dtypes(include=["object"]).columns.tolist()
        if not text_columns:
            return False, "No text-based columns found in the sheet", []

        logger.info(f"[EXCEL] Text columns detected for embedding: {text_columns}")

        # -------- Step 3: Normalize & Combine Rows --------
        df = df.fillna("")

        def clean_value(value: str) -> str:
            """Normalize whitespace."""
            return " ".join(value.split())

        def combine_row(row):
            """Convert entire row into semantically meaningful string."""
            return "\n".join([
                f"{col}: {clean_value(row[col])}"
                for col in text_columns
                if row[col].strip()
            ])

        df["combined_text"] = df.apply(combine_row, axis=1)

        # Filter out fully empty rows
        df = df[df["combined_text"].str.strip() != ""]
        if df.empty:
            return False, "All rows became empty after normalization", []

        # -------- Step 4: Chunking --------
        chunks = []
        total_rows = len(df)

        for idx, text in df["combined_text"].items():
            row_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type="excel")
        
            for chunk in row_chunks:
                if isinstance(chunk, dict):
                    normalized_text = chunk.get("text", "")
                elif isinstance(chunk, str):
                    normalized_text = chunk
                else:
                    normalized_text = ""

                if not isinstance(normalized_text, str) or not normalized_text.strip():
                    continue  # 
                chunks.append({
                    "text": normalized_text,
                    "source_type": "excel",
                    "attributes": extract_attributes_from_text(normalized_text),
                    "metadata": {
                        "row_index": int(idx + 1),
                        "filename": os.path.basename(file_path),
                        "total_rows": total_rows
                    }
                })

        chunks = [c for c in chunks if c["text"] and isinstance(c["text"], str)]

        if not chunks:
            return False, "No valid chunks (all were empty or invalid)", []


        logger.info(f"[EXCEL] Chunking complete, generated {len(chunks)} chunks")

        return True, "Excel processed successfully", chunks

    except Exception as e:
        logger.exception("[EXCEL] Unhandled processing error")
        return False, f"Unexpected error during Excel processing: {str(e)}", []

def parse_urls_input(urls_input: str) -> List[str]:
    """Parse URLs from various input formats (JSON array, comma-separated, or single URL)."""
    if not urls_input or not urls_input.strip():
        return []
    
    urls_input = urls_input.strip()
    
    # Try to parse as JSON array first
    try:
        if urls_input.startswith('[') and urls_input.endswith(']'):
            urls_list = json.loads(urls_input)
            if isinstance(urls_list, list):
                return [url.strip() for url in urls_list if url.strip()]
    except json.JSONDecodeError:
        pass
    
    # Try comma-separated URLs
    if ',' in urls_input:
        return [url.strip() for url in urls_input.split(',') if url.strip()]
    
    # Single URL
    return [urls_input] if urls_input else []

def validate_url(url: str) -> bool:
    """Validate if a URL is properly formatted."""
    try:
        return url.startswith(('http://', 'https://')) and len(url) > 8
    except:
        return False
        
# def crawl_webpage(scrap_url: str, max_crawl_pages: str, max_crawl_depth: str, dynamic_wait: str) -> tuple:
#     """Crawl a webpage and extract clean text using Firecrawl."""
#     try:
#         CRAWLER_API_URL = "https://api.firecrawl.dev/v0/scrape"
#         FIRECRAWL_API_KEY = os.getenv("FIRECRAWL_API_KEY")
        
#         if not FIRECRAWL_API_KEY:
#             logger.error("FIRECRAWL_API_KEY environment variable not set")
#             return False, "FIRECRAWL_API_KEY not configured. Please set the environment variable."
        
#         logger.info(f"Crawling URL: {scrap_url} with Firecrawl API")
        
#         headers = {
#             "Authorization": f"Bearer {FIRECRAWL_API_KEY}",
#             "Content-Type": "application/json"
#         }
        
#         payload = {
#             "url": scrap_url,
#             "crawlerOptions": {
#                 "maxPages": int(max_crawl_pages or 100),
#                 "maxDepth": int(max_crawl_depth or 100),
#                 "waitBetweenRequests": int(dynamic_wait or 1) * 1000,
#                 "onlyMainContent": False,   # <-- get the full HTML/text
#                 "excludes": ["*/login", "*/signup", "*.pdf"],
#                 "returnOnlyUrls": False
#             }
#         }

        
#         logger.debug(f"Firecrawl payload: {payload}")
        
#         response = requests.post(CRAWLER_API_URL, json=payload, headers=headers, timeout=300)
#         logger.info(f"Firecrawl response status: {response.status_code}")
        
#         if response.status_code == 200:
#             try:
#                 data = response.json()
#                 logger.info(f"Firecrawl responseDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDDD-> {data}")
#                 logger.debug(f"Firecrawl response data keys: {data.keys() if isinstance(data, dict) else 'Not a dict'}")
                
#                 if data.get("success") and "data" in data and "content" in data["data"]:
#                     extracted_text = data["data"]["content"]
#                     if not extracted_text or not extracted_text.strip():
#                         logger.warning(f"No meaningful content extracted from {scrap_url}")
#                         return False, "No meaningful content extracted from webpage"
                    
#                     logger.info(f"Successfully extracted {len(extracted_text)} characters from {scrap_url}")
#                     return True, extracted_text
#                 else:
#                     error_msg = data.get('error', 'Unknown error in Firecrawl response')
#                     logger.error(f"Unexpected Firecrawl response structure: {data}")
#                     return False, f"Failed to extract data: {error_msg}"
#             except ValueError as e:
#                 logger.error(f"Error parsing Firecrawl response: {str(e)}")
#                 return False, f"Invalid response format: {str(e)}"
#         else:
#             error_text = response.text[:500]  # Limit error text length
#             logger.error(f"Firecrawl API error: {response.status_code} - {error_text}")
#             return False, f"API request failed (Status {response.status_code}): {error_text}"
#     except requests.exceptions.Timeout:
#         logger.error(f"Timeout while crawling {scrap_url}")
#         return False, "Request timeout - webpage took too long to respond"
#     except requests.exceptions.ConnectionError:
#         logger.error(f"Connection error while crawling {scrap_url}")
#         return False, "Connection error - unable to reach webpage"
#     except Exception as e:
#         logger.error(f"Exception occurred during web crawling of {scrap_url}: {str(e)}")
#         return False, f"Unexpected error: {str(e)}"

def classify_firecrawl_error(error_text: str) -> dict:
    error_text = error_text.lower()
    if "insufficient credits" in error_text:
        return {"error_type": "credits_exhausted"}
    elif "rate limit" in error_text:
        return {"error_type": "rate_limited"}
    elif "timeout" in error_text:
        return {"error_type": "timeout"}
    elif "forbidden" in error_text or "403" in error_text:
        return {"error_type": "blocked"}
    return {"error_type": "unknown"}

def crawl_webpage(
    scrap_url: str,
    max_crawl_pages: str = "",
    max_crawl_depth: str = "",
    dynamic_wait: str = ""
) -> Tuple[bool, str, dict]:
    """
    Uses Firecrawl v2 API to crawl websites.
    
    Returns:
        Tuple[bool, str, dict]: (success, content_or_error, stats)
        stats contains: pages_crawled, credits_used, crawl_mode
    """
    FIRECRAWL_API_KEY = os.getenv("FIRECRAWL_API_KEY")
  
    if not FIRECRAWL_API_KEY:
        logger.error("FIRECRAWL_API_KEY not set")
        return False, "FIRECRAWL_API_KEY not configured", {}

    original_url = _normalize_url_for_crawl(scrap_url)
    scrap_url, redirected = _resolve_redirect_url(scrap_url)
    # Guardrail: avoid seeding domain crawl from content hubs (e.g. /blog/) when root URL was provided.
    if redirected:
        orig_path = urlsplit(original_url).path or "/"
        new_path = (urlsplit(scrap_url).path or "/").lower()
        if orig_path == "/" and new_path.startswith("/blog"):
            logger.warning(f"Redirect seed points to blog path; preserving original root seed: {original_url}")
            scrap_url = original_url
    if redirected:
        logger.info(f"Resolved redirect to final URL: {scrap_url}")

    headers = {
        "Authorization": f"Bearer {FIRECRAWL_API_KEY}",
        "Content-Type": "application/json"
    }

    # Parse and validate inputs
    try:
        depth = max(0, int(max_crawl_depth or 10))
        pages = max(1, int(max_crawl_pages or 25))
        wait_ms = max(0, int(dynamic_wait or 2)) * 1000

        # Enforce system limits
        MAX_ALLOWED_PAGES = 100
        MAX_ALLOWED_DEPTH = 60
        MAX_WAIT_MS = 30000

        pages = min(pages, MAX_ALLOWED_PAGES)
        depth = min(depth, MAX_ALLOWED_DEPTH)
        wait_ms = min(wait_ms, MAX_WAIT_MS)

    except (ValueError, TypeError) as e:
        logger.warning(f"Invalid crawl params: {e}, using defaults")
        depth = 20
        pages = 45
        wait_ms = 2000
    
    analysis = _pre_analyze_crawl_target(scrap_url, headers)
    mode = analysis.get("selected_mode", "STANDARD")
    mode_flags = analysis.get("flags", [])

    logger.info("="*60)
    logger.info(f"CRAWL REQUEST → {scrap_url}")
    logger.info(f"   Pages: {pages} | Depth: {depth} | Wait: {wait_ms}ms")
    logger.info(f"   Mode: {'🕷️ MULTI-PAGE CRAWL' if depth > 0 else '📄 SINGLE PAGE'}")
    logger.info(
        "   Auto-Mode: %s | Flags: %s | Domains: %s | Redirects: %s",
        mode,
        ",".join(mode_flags),
        len(analysis.get("related_domains", [])),
        max(0, len(analysis.get("redirect_chain", [])) - 1),
    )
    logger.info("="*60)

    # CASE 1: Single page (depth = 0)
    if depth == 0:
        success, content = _scrape_single_page_v2(scrap_url, headers)
        stats = {"pages_crawled": 1 if success else 0, "credits_used": 1 if success else 0, "crawl_mode": "single"}
        return success, content, stats

    # CASE 2: Multi-page crawl (depth > 0)
    return _crawl_multi_page_v2(scrap_url, pages, depth, wait_ms, headers, analysis)


def _scrape_single_page_v2(url: str, headers: dict) -> Tuple[bool, str]:
    """Single page scrape using Firecrawl v2 API"""
    payload = {
        "url": url,
        "formats": ["markdown"],
        "onlyMainContent": True,
        "onlyCleanContent": True,
        "waitFor": 2000,
        "blockAds": True,
        "proxy": "auto"
    }
    
    try:
        response = requests.post(
            "https://api.firecrawl.dev/v2/scrape",
            json=payload,
            headers=headers,
            timeout=60
        )
        
        if response.status_code == 200:
            data = response.json()
            
            if data.get("success"):
                content = data.get("data", {}).get("markdown", "")
                if content:
                    resolved_url = data.get("data", {}).get("metadata", {}).get("sourceURL", url)
                    if not _is_useful_scraped_content(content, resolved_url):
                        return False, "No meaningful content extracted"
                    content = f"# Source: {resolved_url}\n\n{content}" if resolved_url else content
                    logger.info(f"✅ Scraped {len(content)} chars from {resolved_url}")
                    return True, content
                else:
                    logger.warning(f"⚠️ No content in response")
                    return False, "No content extracted"
            else:
                error = data.get("error", "Unknown error")
                logger.error(f"❌ Scrape failed: {error}")
                return False, error
        else:
            logger.error(f"❌ HTTP {response.status_code}: {response.text[:300]}")
            return False, f"HTTP {response.status_code}"
            
    except requests.Timeout:
        logger.error(f"❌ Timeout scraping {url}")
        return False, "Request timeout"
    except Exception as e:
        logger.error(f"❌ Exception: {str(e)}")
        return False, str(e)


def _crawl_multi_page_v2(
    url: str, 
    pages: int, 
    depth: int, 
    wait_ms: int, 
    headers: dict,
    analysis: Optional[dict] = None,
) -> Tuple[bool, str, dict]:
    """Multi-page crawl using Firecrawl v2 API. Returns (success, content, stats)."""
    SCRAPED_PREVIEW_CHARS = 1200
    def _extract_page_url(page: dict, fallback_url: str = "unknown") -> str:
        raw_url = (
            page.get("url")
            or page.get("metadata", {}).get("sourceURL")
            or page.get("metadata", {}).get("sourceUrl")
            or fallback_url
        )
        return _normalize_url_for_crawl(raw_url) if raw_url != "unknown" else "unknown"

    def _collect_crawled_content(crawled_data: List[dict], seed_url: str) -> Tuple[str, List[str], List[dict]]:
        all_content, crawled_urls, skipped = [], [], []
        for page in crawled_data or []:
            markdown = page.get("markdown", "")
            page_url = _extract_page_url(page, seed_url)
            if markdown:
                if not _is_useful_scraped_content(markdown, page_url):
                    skipped.append({"url": page_url, "reason": "low_value_or_short_content"})
                    continue
                crawled_urls.append(page_url)
                all_content.append(f"# Source: {page_url}\n\n{markdown}")
                logger.info(
                    "[KB SCRAPED PAGE] url=%s chars=%s preview=%s",
                    page_url,
                    len(markdown),
                    markdown[:SCRAPED_PREVIEW_CHARS].replace("\n", " "),
                )
            else:
                skipped.append({"url": page_url, "reason": "empty_markdown"})
        combined = "\n\n---\n\n".join(all_content)
        return combined, crawled_urls, skipped

    def _priority_sort_key(link: str) -> tuple:
        parsed = urlsplit(link)
        path = parsed.path or "/"
        # Prefer clean, short, non-query URLs first to improve KB signal quality.
        return (
            1 if parsed.query else 0,
            path.count("/"),
            len(path),
            len(link),
        )

    analysis = analysis or {}
    mode = analysis.get("selected_mode", "STANDARD")
    mode_flags = analysis.get("flags", ["DEDUP_STRICT", "REDIRECT_CHAIN_TRACK"])
    related_domains = analysis.get("related_domains", _domain_variants(url))
    stats = {
        "pages_discovered": 0,
        "pages_crawled": 0,
        "credits_used": 0,
        "crawl_mode": "multi",
        "crawl_mode_selected": mode,
        "crawl_mode_flags": mode_flags,
        "allowed_domains": related_domains[:10],
        "crawled_urls": [],
        "skipped_urls": [],
        "redirect_chain": analysis.get("redirect_chain", [url]),
        "duplicates_removed": 0,
    }
    
    # Build v2 payload - CRITICAL CHANGES HERE
    # Keep render wait bounded; very high values (e.g. 15s) dramatically slow throughput.
    render_wait_ms = min(max(wait_ms, 1500), 4000)
    if mode == "JS_RENDERED":
        render_wait_ms = min(max(wait_ms, 3500), 7000)
    crawl_depth = depth
    crawl_limit = pages
    if mode == "DEEP_LINK":
        # Respect caller intent; allow only moderate expansion for deep sites.
        crawl_depth = min(60, max(depth, min(20, depth + max(2, depth // 2))))
        crawl_limit = min(100, max(pages, min(50, pages + max(5, pages // 2))))
    elif mode == "STANDARD":
        crawl_depth = min(depth, 25)

    payload = {
        "url": url,
        "limit": crawl_limit,
        "maxDiscoveryDepth": crawl_depth,
        "crawlEntireDomain": True,
        "allowSubdomains": True,
        "sitemap": "include",
        "ignoreQueryParameters": True,
        "scrapeOptions": {
            "formats": ["markdown"],
            "onlyMainContent": True,
            "onlyCleanContent": True,
            "waitFor": render_wait_ms,
            "blockAds": True,
            "proxy": "auto"
        }
    }
    
    # v2 API DOES NOT support waitBetweenRequests
    # It's handled automatically by Firecrawl
    # if wait_ms > 0:
    #     payload["waitBetweenRequests"] = wait_ms  # ❌ REMOVED - causes 400 error
    
    # Optional: Add exclude patterns if needed
    # Note: v2 uses different syntax - check current docs
    # payload["excludePaths"] = ["*/login", "*/signup"]
    if related_domains:
        payload["allowExternalLinks"] = True if mode == "MULTI_DOMAIN_GRAPH" else False
        payload["allowedDomains"] = related_domains[:12]

    logger.info(
        "[KB CRAWL MODE] mode=%s flags=%s depth=%s limit=%s render_wait=%sms allowed_domains=%s",
        mode,
        ",".join(mode_flags),
        crawl_depth,
        crawl_limit,
        render_wait_ms,
        related_domains[:12],
    )

    try:
        # 1. Start crawl job
        logger.info(f"🕷️ Starting v2 crawl job...")
        response = requests.post(
            "https://api.firecrawl.dev/v2/crawl",
            json=payload,
            headers=headers,
            timeout=120
        )
        if response.status_code == 400 and ("allowedDomains" in payload or "allowExternalLinks" in payload):
            logger.warning("[KB CRAWL MODE] Firecrawl rejected optional domain controls, retrying with base payload")
            payload.pop("allowedDomains", None)
            payload.pop("allowExternalLinks", None)
            response = requests.post(
                "https://api.firecrawl.dev/v2/crawl",
                json=payload,
                headers=headers,
                timeout=120
            )

        if response.status_code != 200:
            logger.error(f"❌ Failed to start crawl: HTTP {response.status_code}")
            logger.error(f"Response: {response.text[:500]}")
            return False, f"Crawl start failed: HTTP {response.status_code}", stats

        job_data = response.json()
        
        # Check for success
        if not job_data.get("success"):
            error = job_data.get("error", "Unknown error")
            logger.error(f"❌ Crawl start failed: {error}")
            return False, f"Crawl start error: {error}", stats
        
        job_id = job_data.get("id")
        if not job_id:
            logger.error(f"❌ No job ID in response: {job_data}")
            return False, "No job ID returned", stats

        logger.info(f"✅ Crawl job started: {job_id}")

        # 2. Poll for completion
        # Dynamic timeout for deeper/larger crawls (prevents false failures while still bounded).
        max_wait_time = min(1800, max(600, 600 + (pages * 8) + (depth * 2)))
        start_time = time.time()
        poll_interval = 4
        max_poll_interval = 20
        attempt = 0
        last_status_data = {}
        zero_progress_attempts = 0

        while time.time() - start_time < max_wait_time:
            attempt += 1
            time.sleep(poll_interval)
            
            try:
                status_response = requests.get(
                    f"https://api.firecrawl.dev/v2/crawl/{job_id}",
                    headers=headers,
                    timeout=30
                )
                
                if status_response.status_code != 200:
                    logger.warning(f"⚠️ Status check failed: HTTP {status_response.status_code}")
                    poll_interval = min(poll_interval * 1.5, max_poll_interval)
                    continue

                status_data = status_response.json()
                last_status_data = status_data
                
                # v2 uses different status structure
                status = status_data.get("status")
                total = status_data.get("total", 0)
                completed = status_data.get("completed", 0)
                credits_used = status_data.get("creditsUsed", 0)
                stats["pages_discovered"] = total or stats["pages_discovered"]
                stats["pages_crawled"] = max(stats["pages_crawled"], completed or 0)
                stats["credits_used"] = max(stats["credits_used"], credits_used or 0)
                
                logger.info(f"📊 Attempt {attempt}: {status} | Pages: {completed}/{total} | Credits: {credits_used}")
                if status in ["scraping", "pending", "processing"] and (completed or 0) == 0 and (credits_used or 0) == 0:
                    zero_progress_attempts += 1
                else:
                    zero_progress_attempts = 0

                if zero_progress_attempts >= 12:
                    logger.warning(
                        "[KB CRAWL MODE] No crawl progress after %s polls; falling back to single-page scrape for %s",
                        zero_progress_attempts, url
                    )
                    single_ok, single_content = _scrape_single_page_v2(url, headers)
                    if single_ok and single_content:
                        stats["crawl_mode"] = "single_fallback"
                        stats["pages_discovered"] = max(stats["pages_discovered"], 1)
                        stats["pages_crawled"] = max(stats["pages_crawled"], 1)
                        stats["credits_used"] = max(stats["credits_used"], 1)
                        stats["crawled_urls"] = [url]
                        return True, single_content, stats

                # 3. Handle completion
                if status == "completed":
                    crawled_data = status_data.get("data", [])
                    stats["pages_crawled"] = len(crawled_data)
                    stats["credits_used"] = credits_used
                    stats["total_requested"] = total
                    logger.info(f"✅ Crawl completed! {len(crawled_data)} pages | {credits_used} credits used")
                    
                    if not crawled_data:
                        logger.warning("⚠️ No pages in completed crawl")
                        return False, "Crawl completed but no pages returned", stats
                    
                    combined, crawled_urls, skipped = _collect_crawled_content(crawled_data, url)
                    unique_crawled_urls = []
                    seen_urls = set()
                    for u in crawled_urls:
                        if u not in seen_urls:
                            unique_crawled_urls.append(u)
                            seen_urls.add(u)
                    stats["duplicates_removed"] = max(0, len(crawled_urls) - len(unique_crawled_urls))
                    stats["skipped_urls"].extend(skipped)
                    stats["crawled_urls"] = unique_crawled_urls
                    if unique_crawled_urls:
                        first_url = unique_crawled_urls[0]
                        if first_url != url:
                            stats["redirect_chain"] = [url, first_url]
                    
                    if not combined.strip():
                        return False, "No content extracted from crawled pages", stats

                    # Dynamic coverage recovery:
                    # Firecrawl may miss important internal pages in some runs.
                    # We use pre-analysis links from the same domain graph and recover uncrawled URLs safely.
                    candidate_links = list(analysis.get("internal_links_sample", []) or [])
                    # Dynamically harvest URLs from sitemap pages when available.
                    sitemap_links = []
                    for page in crawled_data or []:
                        page_url = _extract_page_url(page, url).lower()
                        if not page_url.endswith("sitemap.xml") and "/sitemap" not in page_url:
                            continue
                        page_markdown = page.get("markdown", "") or ""
                        try:
                            extracted = re.findall(r"https?://[^\s<>\")']+", page_markdown)
                        except re.error:
                            extracted = []
                        sitemap_links.extend([_normalize_url_for_crawl(u) for u in extracted if u.startswith("http")])
                    if sitemap_links:
                        candidate_links.extend(sitemap_links)
                        logger.info("[KB COVERAGE RECOVERY] harvested_from_sitemap=%s", len(set(sitemap_links)))
                    crawled_set = set(unique_crawled_urls)
                    skipped_set = {x.get("url") for x in stats["skipped_urls"] if x.get("url")}
                    domain_allow = set(related_domains or [])
                    recover_candidates = []
                    for link in candidate_links:
                        host = (urlsplit(link).netloc or "").lower()
                        if domain_allow and host not in domain_allow:
                            continue
                        if link in crawled_set or link in skipped_set:
                            continue
                        if not link.startswith("http://") and not link.startswith("https://"):
                            continue
                        recover_candidates.append(link)

                    recover_candidates = sorted(set(recover_candidates), key=_priority_sort_key)
                    # Allow bounded post-crawl recovery even when Firecrawl already hit the limit.
                    # This improves coverage for missed but relevant internal links discovered dynamically.
                    extra_recovery_cap = min(20, max(5, pages // 2))
                    recover_budget = min(len(recover_candidates), extra_recovery_cap)
                    recovered_count = 0
                    if recover_budget > 0 and recover_candidates:
                        logger.info(
                            "[KB COVERAGE RECOVERY] attempting=%s candidates=%s crawled=%s limit=%s",
                            recover_budget, len(recover_candidates), len(unique_crawled_urls), crawl_limit
                        )
                    for link in recover_candidates[:recover_budget]:
                        single_ok, single_content = _scrape_single_page_v2(link, headers)
                        if not single_ok or not single_content:
                            stats["skipped_urls"].append({"url": link, "reason": "recovery_scrape_failed_or_low_value"})
                            continue
                        # _scrape_single_page_v2 already returns '# Source: <url>\\n\\n...'
                        combined += "\n\n---\n\n" + single_content
                        unique_crawled_urls.append(link)
                        crawled_set.add(link)
                        recovered_count += 1
                        logger.info(
                            "[KB SCRAPED RECOVERY PAGE] url=%s chars=%s preview=%s",
                            link,
                            len(single_content),
                            single_content[:SCRAPED_PREVIEW_CHARS].replace("\n", " "),
                        )
                    if recovered_count:
                        stats["pages_crawled"] = len(unique_crawled_urls)
                        stats["pages_discovered"] = max(stats["pages_discovered"], len(unique_crawled_urls))
                        logger.info(
                            "[KB COVERAGE RECOVERY] recovered=%s final_crawled=%s",
                            recovered_count, len(unique_crawled_urls)
                        )
                    
                    logger.info(f"✅ Total content: {len(combined)} chars from {len(crawled_data)} pages")
                    return True, combined, stats

                # 4. Handle failure
                elif status in ["failed", "cancelled"]:
                    error_msg = status_data.get("error", "Unknown error")
                    stats["pages_crawled"] = completed
                    stats["credits_used"] = credits_used
                    logger.error(f"❌ Crawl {status}: {error_msg}")
                    logger.error(f"Partial results: {completed}/{total} pages")
                    return False, f"Crawl {status}: {error_msg}", stats

                # 5. Still running
                elif status in ["scraping", "pending", "processing"]:
                    poll_interval = min(poll_interval * 1.2, max_poll_interval)
                
                else:
                    logger.warning(f"⚠️ Unknown status: {status}")

            except requests.Timeout:
                logger.warning(f"⚠️ Status check timeout")
                poll_interval = min(poll_interval * 1.5, max_poll_interval)
                continue
            except Exception as poll_error:
                logger.warning(f"⚠️ Polling error: {poll_error}")
                continue

        # Timeout reached
        elapsed = time.time() - start_time
        logger.error(f"❌ Crawl timeout after {elapsed:.0f}s")
        # Timeout salvage: if Firecrawl already returned partial page data, use it.
        partial_data = (last_status_data or {}).get("data", []) or []
        if partial_data:
            combined, crawled_urls, skipped = _collect_crawled_content(partial_data, url)
            unique_crawled_urls = []
            seen_urls = set()
            for u in crawled_urls:
                if u not in seen_urls:
                    unique_crawled_urls.append(u)
                    seen_urls.add(u)
            stats["duplicates_removed"] = max(0, len(crawled_urls) - len(unique_crawled_urls))
            stats["pages_crawled"] = max(stats["pages_crawled"], len(partial_data))
            stats["crawled_urls"] = unique_crawled_urls
            stats["skipped_urls"].extend(skipped)
            if unique_crawled_urls:
                first_url = unique_crawled_urls[0]
                if first_url != url:
                    stats["redirect_chain"] = [url, first_url]
            if combined.strip():
                logger.warning(
                    "⚠️ Crawl timed out but returning partial content: pages=%s chars=%s",
                    len(partial_data), len(combined)
                )
                return True, combined, stats
        return False, f"Crawl did not complete within {max_wait_time}s", stats

    except requests.Timeout:
        logger.error(f"❌ Request timeout starting crawl")
        return False, "Request timeout", stats
    except Exception as e:
        logger.error(f"❌ Crawl exception: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return False, f"Crawl error: {str(e)}", stats



# # === Unchanged helper (your original single-page logic) ===
# def _single_page_scrape(scrap_url: str, headers: dict) -> Tuple[bool, str]:
#     payload = {
#         "url": scrap_url,
#         "pageOptions": {"onlyMainContent": True}
#     }
#     try:
#         response = requests.post(
#             "https://api.firecrawl.dev/v0/scrape",
#             json=payload,
#             headers=headers,
#             timeout=120
#         )
#         if response.status_code == 200:
#             data = response.json()
#             if data.get("success") and data.get("data", {}).get("content"):
#                 text = data["data"]["content"]
#                 logger.info(f"Single-page /scrape SUCCESS → {len(text)} chars from {scrap_url}")
#                 return True, text
#             else:
#                 logger.warning(f"/scrape returned success=False: {data}")
#                 return False, data.get("error", "No content")
#         else:
#             logger.warning(f"/scrape failed {response.status_code}: {response.text[:300]}")
#             return False, f"HTTP {response.status_code}"
#     except Exception as e:
#         logger.error(f"Single-page scrape exception: {e}")
#         return False, str(e)

def process_multiple_urls(urls: List[str], knowledge_base_name: str, chunk_size: int, chunk_overlap: int, max_crawl_pages: str, max_crawl_depth: str, dynamic_wait: str) -> tuple:
    """Process multiple URLs and return combined chunks and processing results."""
    all_chunks = []
    processed_urls = []
    failed_urls = []
    seen_resolved_urls = set()
    
    for url in urls:
        if not validate_url(url):
            logger.warning(f"Invalid URL format: {url}")
            failed_urls.append({"url": url, "error": "Invalid URL format"})
            continue

        resolved_url, redirected = _resolve_redirect_url(url)
        if resolved_url in seen_resolved_urls:
            logger.info(f"Skipping duplicate resolved URL: {resolved_url} (from {url})")
            continue
        seen_resolved_urls.add(resolved_url)
        
        try:
            logger.info(f"Processing URL: {url} -> {resolved_url}" if redirected else f"Processing URL: {url}")
            success, extracted_text, crawl_stats = crawl_webpage(resolved_url, max_crawl_pages, max_crawl_depth, dynamic_wait)
            
            if success:
                url_chunks = chunk_text(extracted_text, chunk_size, chunk_overlap)
                all_chunks.extend(url_chunks)
                processed_urls.append(resolved_url)
                create_website_chunks_directory(knowledge_base_name, resolved_url)
                logger.info(f"Successfully processed URL: {resolved_url}, extracted {len(url_chunks)} chunks")
            else:
                logger.error(f"Failed to process URL: {resolved_url}, error: {extracted_text}")
                failed_urls.append({"url": resolved_url, "error": extracted_text})
        except Exception as e:
            logger.error(f"Exception while processing URL {resolved_url}: {str(e)}")
            failed_urls.append({"url": resolved_url, "error": str(e)})
    
    return all_chunks, processed_urls, failed_urls

def chunk_text(text: str, chunk_size: int, overlap: int, source_type: str) -> List[dict]:
    """Split text into chunks, returning a list of dictionaries with text and source_type."""
    if not text.strip():
        return []
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()
        if chunk:  # Only add non-empty chunks
            chunks.append({"text": chunk, "source_type": source_type})
        start = end - overlap
    logger.debug(f"Created {len(chunks)} chunks from {source_type} source")
    return chunks


def chunk_web_content_by_source(
    combined_text: str,
    chunk_size: int,
    overlap: int,
    crawl_meta: Optional[dict] = None
) -> List[dict]:
    """
    Split merged crawl content back into page-level segments using '# Source: <url>'
    markers, then chunk each page while preserving source_url metadata.
    """
    if not combined_text or not combined_text.strip():
        return []

    parts = re.split(r"(?m)^# Source:\s+", combined_text)
    chunks: List[dict] = []
    crawl_meta = crawl_meta or {}
    mode = crawl_meta.get("crawl_mode_selected", "STANDARD")

    for part in parts:
        part = part.strip()
        if not part:
            continue
        lines = part.splitlines()
        source_url = lines[0].strip() if lines else "unknown"
        page_body = "\n".join(lines[1:]).strip()
        if not page_body:
            continue

        start = 0
        while start < len(page_body):
            end = start + chunk_size
            piece = page_body[start:end].strip()
            if piece:
                chunks.append({
                    "text": piece,
                    "source_type": "web",
                    "source_url": source_url,
                    "source_domain": (urlsplit(source_url).netloc.lower() if source_url and source_url != "unknown" else "unknown"),
                    "crawl_mode": mode,
                })
            start = end - overlap

    logger.debug(f"Created {len(chunks)} chunks from web source with source_url metadata")
    return chunks
    
# @knowledge_base_blueprint.route("/register", methods=["POST"])
# @jwt_required()
# def create_knowledge_base():
#     """Register a new knowledge base with multiple files, web content, and raw text."""
#     try:
#         logger.info("=== KNOWLEDGE BASE CREATION REQUEST ===")
#         logger.debug(f"Request headers: {dict(request.headers)}")
#         logger.debug(f"Content-Type: {request.content_type}")

#         # JWT validation
#         identity = get_jwt_identity()
#         claims = get_jwt()
#         tenant_id = claims.get("tenant_id")
#         if not tenant_id:
#             logger.error("Tenant ID not found in token")
#             return jsonify({"status": "error", "message": "Tenant ID not found in token"}), 401

#         # Content-Type check
#         if not request.content_type or 'multipart/form-data' not in request.content_type:
#             logger.error("Invalid Content-Type: Expected multipart/form-data")
#             return jsonify({"status": "error", "message": "Content-Type must be multipart/form-data"}), 400

#         # Form data and files
#         try:
#             data = request.form
#             pdf_files = request.files.getlist("pdf_files")
#             excel_files = request.files.getlist("excel_files")
#             docx_files = request.files.getlist("docx_files")
#             txt_files = request.files.getlist("txt_files")
            
#             files = pdf_files + excel_files + docx_files + txt_files

#             logger.info(f"Form fields received: {list(data.keys())}")
#             logger.info(f"Files received: {[f.filename for f in files if f.filename]}")
#         except Exception as e:
#             logger.error(f"Error accessing form data or files: {str(e)}")
#             return jsonify({"status": "error", "message": "Invalid or corrupted form data"}), 400

#         # Required field check
#         required_fields = ["knowledge_base_name"]
#         if not all(field in data and data[field].strip() for field in required_fields):
#             logger.warning("Missing required field: knowledge_base_name")
#             return jsonify({"status": "error", "message": "Missing required field: knowledge_base_name"}), 400

#         # Extract parameters
#         knowledge_base_name = data["knowledge_base_name"].strip()
#         description = data.get("description", "").strip()
#         scrap_urls = [data[key].strip() for key in data if key.startswith("scrap_urls") and data[key].strip()]
#         raw_text = data.get("raw_text", "").strip()
#         provider = data.get("provider", "").strip()
#         embedding_model = data.get("embeddingModel", "").strip()
#         secret_key = data.get("secretKey", "").strip()
#         chunk_size = int(data.get("chunk_size", 300))
#         chunk_overlap = int(data.get("chunk_overlap", 50))
#         max_crawl_pages = data.get("max_crawl_pages", "")
#         max_crawl_depth = data.get("max_crawl_depth", "")
#         dynamic_wait = data.get("dynamic_wait", "")

#         logger.info(f"Parameters: name={knowledge_base_name}, urls={scrap_urls}, raw_text_length={len(raw_text)}, files={len(files)}")

#         # Validate chunk parameters
#         if chunk_size < 100 or chunk_size > 10000:
#             logger.warning(f"Invalid chunk_size: {chunk_size}")
#             return jsonify({"status": "error", "message": "Chunk size must be between 100 and 10000"}), 400
#         if chunk_overlap < 0:
#             logger.warning(f"Invalid chunk_overlap: {chunk_overlap}")
#             return jsonify({"status": "error", "message": "Chunk overlap must be non-negative"}), 400

#         # Validate embedding model
#         try:
#             embedding_config = get_embedding_model(provider, embedding_model, secret_key)
#         except ValueError as e:
#             logger.error(f"Embedding model initialization failed: {str(e)}")
#             return jsonify({"status": "error", "message": str(e)}), 400

#         chunks = []
#         processed_files = []
#         processed_urls = []
#         failed_urls = []
#         upload_dir = "Uploads"
#         os.makedirs(upload_dir, exist_ok=True)

#         # Allowed file types
#         allowed_extensions = {'.pdf', '.txt', '.xls', '.xlsx'}
#         if DOCX_AVAILABLE:
#             allowed_extensions.add('.docx')
#         invalid_files = [file.filename for file in files if file.filename and os.path.splitext(file.filename)[1].lower() not in allowed_extensions]
#         if invalid_files:
#             logger.warning(f"Invalid file types uploaded: {invalid_files}")
#             return jsonify({"status": "error", "message": f"Only PDF, DOCX, TXT, XLS, XLSX files are allowed: {invalid_files}"}), 400

#         # Process uploaded files
#         for file in files:
#             if not file.filename:
#                 logger.warning("Skipping empty filename")
#                 continue

#             file_ext = os.path.splitext(file.filename)[1].lower()
#             file_name_without_ext = os.path.splitext(file.filename)[0]
#             file_path = os.path.join(upload_dir, file.filename)
#             print("ddddddddddddddddddddddddddddddddd")
#             # Save file
#             try:
#                 file.save(file_path)
#                 logger.info(f"Saved file: {file.filename}")
#             except Exception as e:
#                 logger.error(f"Failed to save file {file.filename}: {str(e)}")
#                 return jsonify({"status": "error", "message": f"Failed to save file {file.filename}: {str(e)}"}), 500

#             try:
#                 # if file_ext in ['.xls', '.xlsx']:
#                 #     df = pd.read_excel(file_path)
#                 #     expected_cols = ['item_code', 'material_description', 'supplier', 'supplier_email', 'notes']
#                 #     if not all(col in df.columns for col in expected_cols):
#                 #         os.remove(file_path)
#                 #         return jsonify({
#                 #             "status": "error",
#                 #             "message": f"Excel {file.filename} must have columns: {expected_cols}"
#                 #         }), 400

#                 #     session = next(db_session())
#                 #     try:
#                 #         for _, row in df.iterrows():
#                 #             supplier = SupplierDetails(
#                 #                 item_code=str(row['item_code']).strip(),
#                 #                 material_description=str(row.get('material_description', '')).strip(),
#                 #                 supplier=str(row['supplier']).strip(),
#                 #                 supplier_email=str(row['supplier_email']).strip(),
#                 #                 notes=str(row.get('notes', '')).strip()
#                 #             )
#                 #             session.add(supplier)

#                 #             row_text = (
#                 #                 f"Item Code: {row['item_code']}, Material: {row.get('material_description','')}, "
#                 #                 f"Supplier: {row['supplier']}, Email: {row['supplier_email']}, Notes: {row.get('notes','')}"
#                 #             )
#                 #             row_chunks = chunk_text(row_text, chunk_size, chunk_overlap, source_type='excel')
#                 #             chunks.extend(row_chunks)

#                 #         session.commit()
#                 #         processed_files.append(file.filename)
#                 #         logger.info(f"Inserted {len(df)} rows from Excel {file.filename}, created {len(row_chunks)} chunks")
#                 #     except Exception as e:
#                 #         session.rollback()
#                 #         raise
#                 #     finally:
#                 #         session.close()
#                 if file_ext in ['.xls', '.xlsx']:
#                     logger.info(f"[REGISTER] Processing Excel file → {file.filename}")

#                     success, message, excel_chunks = process_excel_file(
#                         file_path=file_path,
#                         chunk_size=chunk_size,
#                         chunk_overlap=chunk_overlap,
                        
#                     )

#                     if not success:
#                         logger.error(f"[EXCEL] {file.filename} failed → {message}")
#                         return jsonify({"status": "error", "message": message}), 400

#                     chunks.extend(excel_chunks)
#                     processed_files.append(file.filename)
#                     media_type = "excel"
#                     logger.info(f"[EXCEL] Completed processing for {file.filename} → {len(excel_chunks)} chunks")

#                 elif file_ext == '.pdf':
#                     output_dir = create_directory_structure(knowledge_base_name, file_name_without_ext)
#                     success, message, file_chunks = process_pdf(file_path, output_dir, knowledge_base_name, chunk_size, chunk_overlap)
#                     if not success:
#                         logger.error(f"Failed to process PDF {file.filename}: {message}")
#                         raise Exception(message)
#                     chunks.extend([{"text": chunk, "source_type": "pdf"} for chunk in file_chunks])
#                     processed_files.append(file.filename)
#                     logger.info(f"Processed PDF file: {file.filename}, created {len(file_chunks)} chunks")

#                 elif file_ext == '.docx':
#                     if not DOCX_AVAILABLE:
#                         raise Exception("python-docx not installed")
#                     text = process_docx_file(file_path)
#                     file_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type='docx')
#                     chunks.extend(file_chunks)
#                     processed_files.append(file.filename)
#                     logger.info(f"Processed DOCX file: {file.filename}, created {len(file_chunks)} chunks")

#                 elif file_ext == '.txt':
#                     text = process_txt_file(file_path)
#                     file_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type='txt')
#                     chunks.extend(file_chunks)
#                     processed_files.append(file.filename)
#                     logger.info(f"Processed TXT file: {file.filename}, created {len(file_chunks)} chunks")

#             except Exception as e:
#                 logger.error(f"Failed to process file {file.filename}: {str(e)}")
#                 return jsonify({"status": "error", "message": f"Failed to process file {file.filename}: {str(e)}"}), 500
#             finally:
#                 if os.path.exists(file_path):
#                     os.remove(file_path)
#                     logger.debug(f"Cleaned up file: {file_path}")

#         # Process web-scraped content
#         if scrap_urls:
#             with ThreadPoolExecutor(max_workers=min(len(scrap_urls), 5)) as executor:
#                 future_to_url = {
#                     executor.submit(crawl_webpage, url, max_crawl_pages, max_crawl_depth, dynamic_wait): url
#                     for url in scrap_urls
#                 }
#                 for future in as_completed(future_to_url):
#                     url = future_to_url[future]
#                     try:
#                         success, extracted_text = future.result()
#                         if success:
#                             website_chunks = chunk_text(extracted_text, chunk_size, chunk_overlap, source_type='web')
#                             chunks.extend(website_chunks)
#                             processed_urls.append(url)
#                             create_website_chunks_directory(knowledge_base_name, url)
#                             logger.info(f"Processed web content from {url}, created {len(website_chunks)} chunks")
#                         else:
#                             logger.warning(f"Web crawling failed for {url}: {extracted_text}")
#                             failed_urls.append({"url": url, "error": extracted_text})
#                     except Exception as e:
#                         logger.error(f"Exception processing {url}: {str(e)}")
#                         failed_urls.append({"url": url, "error": str(e)})

#         # Process raw text
#         if raw_text:
#             text_chunks = chunk_text(raw_text, chunk_size, chunk_overlap, source_type='raw')
#             chunks.extend(text_chunks)
#             create_raw_text_chunks_directory(knowledge_base_name)
#             logger.info(f"Processed raw text, created {len(text_chunks)} chunks")

#         # Final content check
#         has_content = bool(chunks) or bool(processed_files) or bool(processed_urls) or bool(raw_text)
#         if not has_content:
#             debug_info = {
#                 "files_received": len(files),
#                 "valid_files": len(processed_files),
#                 "urls_received": len(scrap_urls),
#                 "urls_processed": len(processed_urls),
#                 "failed_urls": failed_urls,
#                 "raw_text_length": len(raw_text)
#             }
#             logger.warning(f"No content provided for knowledge base. Debug info: {debug_info}")
#             return jsonify({
#                 "status": "error",
#                 "message": "No content provided to create knowledge base. Please upload PDF/DOCX/TXT/XLS/XLSX files, raw text, or provide scrap URLs.",
#                 "debug_info": debug_info
#             }), 400

#         collection_name = f"tenant_{tenant_id}_{knowledge_base_name.lower().replace(' ', '_')}"
#         try:
#             create_qdrant_collection(collection_name, embedding_config["vector_size"])
#             insert_chunks_to_qdrant(collection_name, chunks, embedding_config)
#         except ValueError as e:
#             logger.error(f"Qdrant dimension error: {str(e)}")
#             return jsonify({"status": "error", "message": str(e)}), 400
#         except RuntimeError as e:
#             logger.error(f"Qdrant operation failed: {str(e)}")
#             return jsonify({"status": "error", "message": str(e)}), 500

        
        
#         # --- Generate KB summary AFTER inserting embeddings ---
#         from utils import generate_kb_summary_from_chunks

#         kb_summary = ""
#         try:
#             kb_summary = generate_kb_summary_from_chunks(
#                 kb_name=collection_name,
#                 qdrant_client=qdrant,   
#                 limit=6
#             )
#         except Exception as e:
#             logger.error(f"[KB SUMMARY] Unexpected error: {e}. Continuing without summary.")
#             kb_summary = ""

#         if kb_summary and len(kb_summary) > 5000:
#             LOGGER.warning("[KB SUMMARY] Trimming summary to 5000 chars")
#             kb_summary = kb_summary[:5000]

#         if not kb_summary:
#             logger.info("[KB SUMMARY] No summary generated — likely low content or blank KB.")


#         # Save KnowledgeBase record
#         session = next(db_session())
#         try:
      
#             new_knowledge_base = KnowledgeBase(
#                 tenant_id=tenant_id,
#                 knowledge_base_name=knowledge_base_name,
#                 description=description or None,
#                 upload_pdf=", ".join(processed_files) if processed_files else None,
#                 upload_media=", ".join(processed_files) if processed_files else None,  # store actual excel filenames
#                 media_type=media_type if 'media_type' in locals() else None,
#                 scrap_url=", ".join(processed_urls) if processed_urls else None,
#                 max_crawl_pages=int(max_crawl_pages or 0) or None,
#                 max_crawl_depth=int(max_crawl_depth or 0) or None,
#                 dynamic_wait=int(dynamic_wait or 0) or None,
#                 raw_text=raw_text if raw_text else None,
#                 chunk_size=chunk_size,
#                 chunk_overlap=chunk_overlap,
#                 collection_name=collection_name,
#                 kb_summary=kb_summary or None,
#             )

#             session.add(new_knowledge_base)
#             session.commit()

#             response_data = {
#                 "status": "success",
#                 "message": "KnowledgeBase registered successfully",
#                 "data": {
#                     "knowledge_base_id": new_knowledge_base.knowledge_base_id,
#                     "collection_name": collection_name,
#                     "processed_files": processed_files,
#                     "processed_urls": processed_urls,
#                     "failed_urls": failed_urls,
#                     "total_chunks": len(chunks)
#                 }
#             }
#             logger.info(f"KnowledgeBase created: knowledge_base_id={new_knowledge_base.knowledge_base_id}, collection_name={collection_name}")
#             return jsonify(response_data), 201
#         except Exception as e:
#             session.rollback()
#             logger.error(f"Error during KnowledgeBase creation: {str(e)}")
#             return jsonify({"status": "error", "message": f"Failed to save KnowledgeBase to database: {str(e)}"}), 500
#         finally:
#             session.close()

#     except Exception as e:
#         logger.error(f"Unexpected error in create_knowledge_base: {str(e)}")
#         return jsonify({"status": "error", "message": f"Internal server error: {str(e)}"}), 500

@knowledge_base_blueprint.route("/register", methods=["POST"])
@jwt_required()
def create_knowledge_base():
    """Register a new knowledge base with multiple files, web content, and raw text."""
    try:
        logger.info("=== KNOWLEDGE BASE CREATION REQUEST ===")
        logger.debug(f"Request headers: {dict(request.headers)}")
        logger.debug(f"Content-Type: {request.content_type}")

        # JWT validation
        identity = get_jwt_identity()
        claims = get_jwt()
        tenant_id = claims.get("tenant_id")
        if not tenant_id:
            logger.error("Tenant ID not found in token")
            return jsonify({"status": "error", "message": "Tenant ID not found in token"}), 401

        # Content-Type check
        if not request.content_type or 'multipart/form-data' not in request.content_type:
            logger.error("Invalid Content-Type: Expected multipart/form-data")
            return jsonify({"status": "error", "message": "Content-Type must be multipart/form-data"}), 400

        # Form data and files
        try:
            data = request.form
            pdf_files = request.files.getlist("pdf_files")
            excel_files = request.files.getlist("excel_files")
            docx_files = request.files.getlist("docx_files")
            txt_files = request.files.getlist("txt_files")
            
            files = pdf_files + excel_files + docx_files + txt_files

            logger.info(f"Form fields received: {list(data.keys())}")
            logger.info(f"Files received: {[f.filename for f in files if f.filename]}")
        except Exception as e:
            logger.error(f"Error accessing form data or files: {str(e)}")
            return jsonify({"status": "error", "message": "Invalid or corrupted form data"}), 400

        # Required field check
        required_fields = ["knowledge_base_name"]
        if not all(field in data and data[field].strip() for field in required_fields):
            logger.warning("Missing required field: knowledge_base_name")
            return jsonify({"status": "error", "message": "Missing required field: knowledge_base_name"}), 400

        # Extract parameters
        knowledge_base_name = data["knowledge_base_name"].strip()
        description = data.get("description", "").strip()
        scrap_urls = []
        # Case 1: Single field input (JSON / comma-separated / single)
        if "scrap_urls" in data and data.get("scrap_urls").strip():
            scrap_urls = parse_urls_input(data.get("scrap_urls"))

        # Case 2: Multiple dynamic fields (scrap_urls_1, scrap_urls_2...)
        else:
            scrap_urls = [
                data[key].strip()
                for key in data
                if key.startswith("scrap_urls") and data[key].strip()
            ]

        # Normalize + deduplicate
        scrap_urls = list(set([url.strip().rstrip('/') for url in scrap_urls]))

        logger.info(f"[DEBUG] Final parsed URLs: {scrap_urls}")
        valid_urls = []
        invalid_urls = []

        for url in scrap_urls:
            if validate_url(url):
                valid_urls.append(url)
            else:
                invalid_urls.append(url)

        scrap_urls = valid_urls
        logger.info(f"[DEBUG] Valid URLs: {scrap_urls}")
        logger.warning(f"[DEBUG] Invalid URLs: {invalid_urls}")

        if not scrap_urls:
            logger.warning("[WARNING] No valid URLs found after parsing")
            
        raw_text = data.get("raw_text", "").strip()
        provider = data.get("provider", "").strip()
        embedding_model = data.get("embeddingModel", "").strip()
        secret_key = data.get("secretKey", "").strip()
        chunk_size = int(data.get("chunk_size", 1024))
        chunk_overlap = int(data.get("chunk_overlap", 70))
        max_crawl_pages = data.get("max_crawl_pages", "")
        max_crawl_depth = data.get("max_crawl_depth", "")
        dynamic_wait = data.get("dynamic_wait", "")

        logger.info(f"Parameters: name={knowledge_base_name}, urls={scrap_urls}, raw_text_length={len(raw_text)}, files={len(files)}")

        # Validate chunk parameters
        if chunk_size < 100 or chunk_size > 10000:
            logger.warning(f"Invalid chunk_size: {chunk_size}")
            return jsonify({"status": "error", "message": "Chunk size must be between 100 and 10000"}), 400
        if chunk_overlap < 0:
            logger.warning(f"Invalid chunk_overlap: {chunk_overlap}")
            return jsonify({"status": "error", "message": "Chunk overlap must be non-negative"}), 400

        # Validate embedding model
        try:
            embedding_config = get_embedding_model(provider, embedding_model, secret_key)
        except ValueError as e:
            logger.error(f"Embedding model initialization failed: {str(e)}")
            return jsonify({"status": "error", "message": str(e)}), 400

        chunks = []
        processed_files = []
        processed_urls = []
        failed_urls = []
        upload_dir = "Uploads"
        os.makedirs(upload_dir, exist_ok=True)

        # Allowed file types
        allowed_extensions = {'.pdf', '.txt', '.xls', '.xlsx'}
        if DOCX_AVAILABLE:
            allowed_extensions.add('.docx')
        invalid_files = [file.filename for file in files if file.filename and os.path.splitext(file.filename)[1].lower() not in allowed_extensions]
        if invalid_files:
            logger.warning(f"Invalid file types uploaded: {invalid_files}")
            return jsonify({"status": "error", "message": f"Only PDF, DOCX, TXT, XLS, XLSX files are allowed: {invalid_files}"}), 400

        # Process uploaded files
        for file in files:
            if not file.filename:
                logger.warning("Skipping empty filename")
                continue

            file_ext = os.path.splitext(file.filename)[1].lower()
            file_name_without_ext = os.path.splitext(file.filename)[0]
            file_path = os.path.join(upload_dir, file.filename)
            print("ddddddddddddddddddddddddddddddddd")
            # Save file
            try:
                file.save(file_path)
                logger.info(f"Saved file: {file.filename}")
            except Exception as e:
                logger.error(f"Failed to save file {file.filename}: {str(e)}")
                return jsonify({"status": "error", "message": f"Failed to save file {file.filename}: {str(e)}"}), 500

            try:
                # if file_ext in ['.xls', '.xlsx']:
                #     df = pd.read_excel(file_path)
                #     expected_cols = ['item_code', 'material_description', 'supplier', 'supplier_email', 'notes']
                #     if not all(col in df.columns for col in expected_cols):
                #         os.remove(file_path)
                #         return jsonify({
                #             "status": "error",
                #             "message": f"Excel {file.filename} must have columns: {expected_cols}"
                #         }), 400

                #     session = next(db_session())
                #     try:
                #         for _, row in df.iterrows():
                #             supplier = SupplierDetails(
                #                 item_code=str(row['item_code']).strip(),
                #                 material_description=str(row.get('material_description', '')).strip(),
                #                 supplier=str(row['supplier']).strip(),
                #                 supplier_email=str(row['supplier_email']).strip(),
                #                 notes=str(row.get('notes', '')).strip()
                #             )
                #             session.add(supplier)

                #             row_text = (
                #                 f"Item Code: {row['item_code']}, Material: {row.get('material_description','')}, "
                #                 f"Supplier: {row['supplier']}, Email: {row['supplier_email']}, Notes: {row.get('notes','')}"
                #             )
                #             row_chunks = chunk_text(row_text, chunk_size, chunk_overlap, source_type='excel')
                #             chunks.extend(row_chunks)

                #         session.commit()
                #         processed_files.append(file.filename)
                #         logger.info(f"Inserted {len(df)} rows from Excel {file.filename}, created {len(row_chunks)} chunks")
                #     except Exception as e:
                #         session.rollback()
                #         raise
                #     finally:
                #         session.close()
                if file_ext in ['.xls', '.xlsx']:
                    logger.info(f"[REGISTER] Processing Excel file → {file.filename}")

                    success, message, excel_chunks = process_excel_file(
                        file_path=file_path,
                        chunk_size=chunk_size,
                        chunk_overlap=chunk_overlap,
                        
                    )

                    if not success:
                        logger.error(f"[EXCEL] {file.filename} failed → {message}")
                        return jsonify({"status": "error", "message": message}), 400

                    chunks.extend(excel_chunks)
                    processed_files.append(file.filename)
                    media_type = "excel"
                    logger.info(f"[EXCEL] Completed processing for {file.filename} → {len(excel_chunks)} chunks")

                elif file_ext == '.pdf':
                    output_dir = create_directory_structure(knowledge_base_name, file_name_without_ext)
                    success, message, file_chunks = process_pdf(file_path, output_dir, knowledge_base_name, chunk_size, chunk_overlap)
                    if not success:
                        logger.error(f"Failed to process PDF {file.filename}: {message}")
                        raise Exception(message)
                    chunks.extend([{"text": chunk, "source_type": "pdf"} for chunk in file_chunks])
                    processed_files.append(file.filename)
                    logger.info(f"Processed PDF file: {file.filename}, created {len(file_chunks)} chunks")

                elif file_ext == '.docx':
                    if not DOCX_AVAILABLE:
                        raise Exception("python-docx not installed")
                    text = process_docx_file(file_path)
                    file_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type='docx')
                    chunks.extend(file_chunks)
                    processed_files.append(file.filename)
                    logger.info(f"Processed DOCX file: {file.filename}, created {len(file_chunks)} chunks")

                elif file_ext == '.txt':
                    text = process_txt_file(file_path)
                    file_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type='txt')
                    chunks.extend(file_chunks)
                    processed_files.append(file.filename)
                    logger.info(f"Processed TXT file: {file.filename}, created {len(file_chunks)} chunks")

            except Exception as e:
                logger.error(f"Failed to process file {file.filename}: {str(e)}")
                return jsonify({"status": "error", "message": f"Failed to process file {file.filename}: {str(e)}"}), 500
            finally:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    logger.debug(f"Cleaned up file: {file_path}")

        # Process web-scraped content
        if scrap_urls:
            url_chunks, processed_urls, failed_urls = process_multiple_urls(
                scrap_urls,
                knowledge_base_name,
                chunk_size,
                chunk_overlap,
                max_crawl_pages,
                max_crawl_depth,
                dynamic_wait
            )
            chunks.extend(url_chunks)

            logger.info(f"[URL PROCESSING] Success: {len(processed_urls)}, Failed: {len(failed_urls)}") 

        # Process raw text
        if raw_text:
            text_chunks = chunk_text(raw_text, chunk_size, chunk_overlap, source_type='raw')
            chunks.extend(text_chunks)
            create_raw_text_chunks_directory(knowledge_base_name)
            logger.info(f"Processed raw text, created {len(text_chunks)} chunks")

        # Final content check
        has_content = bool(chunks) or bool(processed_files) or bool(processed_urls) or bool(raw_text)
        if not has_content:
            debug_info = {
                "files_received": len(files),
                "valid_files": len(processed_files),
                "urls_received": len(scrap_urls),
                "urls_processed": len(processed_urls),
                "failed_urls": failed_urls,
                "raw_text_length": len(raw_text)
            }
            logger.warning(f"No content provided for knowledge base. Debug info: {debug_info}")
            return jsonify({
                "status": "error",
                "message": "No content provided to create knowledge base. Please upload PDF/DOCX/TXT/XLS/XLSX files, raw text, or provide scrap URLs.",
                "debug_info": debug_info
            }), 400

       

            

        # Create Qdrant collection
        collection_name = f"tenant_{tenant_id}_{knowledge_base_name.lower().replace(' ', '_')}"
        try:
            create_qdrant_collection(collection_name, embedding_config["vector_size"])
            insert_chunks_to_qdrant(collection_name, chunks, embedding_config)
        except ValueError as e:
            logger.error(f"Qdrant dimension error: {str(e)}")
            return jsonify({"status": "error", "message": str(e)}), 400
        except RuntimeError as e:
            logger.error(f"Qdrant operation failed: {str(e)}")
            return jsonify({"status": "error", "message": str(e)}), 500

        
        
        # --- Generate KB summary AFTER inserting embeddings ---
        from utils import generate_kb_summary_from_chunks

        kb_summary = ""
        try:
            kb_summary = generate_kb_summary_from_chunks(
                kb_name=collection_name,
                qdrant_client=qdrant,   
                limit=6
            )
        except Exception as e:
            logger.error(f"[KB SUMMARY] Unexpected error: {e}. Continuing without summary.")
            kb_summary = ""

        if kb_summary and len(kb_summary) > 5000:
            LOGGER.warning("[KB SUMMARY] Trimming summary to 5000 chars")
            kb_summary = kb_summary[:5000]

        if not kb_summary:
            logger.info("[KB SUMMARY] No summary generated — likely low content or blank KB.")


        # Save KnowledgeBase record
        session = next(db_session())
        try:
      
            new_knowledge_base = KnowledgeBase(
                tenant_id=tenant_id,
                knowledge_base_name=knowledge_base_name,
                description=description or None,
                upload_pdf=", ".join(processed_files) if processed_files else None,
                upload_media=", ".join(processed_files) if processed_files else None,  # store actual excel filenames
                media_type=media_type if 'media_type' in locals() else None,
                scrap_url=", ".join(processed_urls) if processed_urls else None,
                max_crawl_pages=int(max_crawl_pages or 0) or None,
                max_crawl_depth=int(max_crawl_depth or 0) or None,
                dynamic_wait=int(dynamic_wait or 0) or None,
                raw_text=raw_text if raw_text else None,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                collection_name=collection_name,
                kb_summary=kb_summary or None,
            )

            session.add(new_knowledge_base)
            session.commit()

            response_data = {
                "status": "success",
                "message": "KnowledgeBase registered successfully",
                "data": {
                    "knowledge_base_id": new_knowledge_base.knowledge_base_id,
                    "collection_name": collection_name,
                    "processed_files": processed_files,
                    "processed_urls": processed_urls,
                    "failed_urls": failed_urls,
                    "total_chunks": len(chunks)
                }
            }
            logger.info(f"KnowledgeBase created: knowledge_base_id={new_knowledge_base.knowledge_base_id}, collection_name={collection_name}")
            return jsonify(response_data), 201
        except Exception as e:
            session.rollback()
            logger.error(f"Error during KnowledgeBase creation: {str(e)}")
            return jsonify({"status": "error", "message": f"Failed to save KnowledgeBase to database: {str(e)}"}), 500
        finally:
            session.close()

    except Exception as e:
        logger.error(f"Unexpected error in create_knowledge_base: {str(e)}")
        return jsonify({"status": "error", "message": f"Internal server error: {str(e)}"}), 500

@knowledge_base_blueprint.route("/", methods=["GET"])
@jwt_required()
def list_knowledge_bases():
    """List all knowledge bases for the authenticated tenant, ordered by creation date (newest first)."""
    try:
        identity = get_jwt_identity()
        claims = get_jwt()
        tenant_id = claims.get("tenant_id")
        if not tenant_id:
            logger.error("Tenant ID not found in token")
            return jsonify({"status": "error", "message": "Tenant ID not found in token"}), 401

        session = next(db_session())
        try:
            # Order by created_at in descending order
            knowledge_bases = session.query(KnowledgeBase).filter_by(tenant_id=tenant_id, del_flg=False).order_by(KnowledgeBase.created_at.desc()).all()
            knowledge_base_list = [{
                "knowledge_base_id": kb.knowledge_base_id,
                "tenant_id": kb.tenant_id,
                "knowledge_base_name": kb.knowledge_base_name,
                "description": kb.description,
                "knowledge_base_summary": kb.kb_summary,
                "upload_pdf": kb.upload_pdf,
                "scrap_url": kb.scrap_url,
                "max_crawl_pages": kb.max_crawl_pages,
                "max_crawl_depth": kb.max_crawl_depth,
                "dynamic_wait": kb.dynamic_wait,
                "raw_text": kb.raw_text,
                "chunk_size": kb.chunk_size,
                "chunk_overlap": kb.chunk_overlap,
                "collection_name": kb.collection_name,
                "created_at": kb.created_at.isoformat() if kb.created_at else None,
                "updated_at": kb.updated_at.isoformat() if kb.updated_at else None
            } for kb in knowledge_bases]
            logger.info(f"Retrieved {len(knowledge_base_list)} knowledge bases for tenant_id={tenant_id}")
            return jsonify({
                "data": knowledge_base_list,
                "status": "success",
                "message": "Knowledge bases retrieved successfully"
            }), 200
        except Exception as e:
            logger.error(f"Error while retrieving knowledge bases: {str(e)}")
            return jsonify({
                "data": [],
                "status": "error",
                "message": "Failed to retrieve knowledge bases"
            }), 500
        finally:
            session.close()
    except Exception as e:
        logger.error(f"Unexpected error in list_knowledge_bases: {str(e)}")
        return jsonify({
            "data": [],
            "status": "error",
            "message": f"Internal server error: {str(e)}"
        }), 500

@knowledge_base_blueprint.route("/bulk-url-upload", methods=["POST"])
@jwt_required()
def bulk_url_upload():
    """
    Dedicated endpoint for bulk URL processing.
    Accepts a JSON payload with multiple URLs for more efficient processing.
    """
    try:
        identity = get_jwt_identity()
        claims = get_jwt()
        tenant_id = claims.get("tenant_id")
        if not tenant_id:
            logger.error("Tenant ID not found in token")
            return jsonify({"status": "error", "message": "Tenant ID not found in token"}), 401

        # Validate Content-Type for JSON
        if not request.content_type or 'application/json' not in request.content_type:
            logger.error("Invalid Content-Type: Expected application/json")
            return jsonify({"status": "error", "message": "Content-Type must be application/json"}), 400

        try:
            data = request.get_json()
            if not data:
                return jsonify({"status": "error", "message": "No JSON data provided"}), 400
        except Exception as e:
            logger.error(f"Error parsing JSON data: {str(e)}")
            return jsonify({"status": "error", "message": "Invalid JSON data"}), 400

        # Validate required fields
        if not data.get("knowledge_base_name") or not data.get("knowledge_base_name").strip():
            return jsonify({"status": "error", "message": "Missing required field: knowledge_base_name"}), 400

        if not data.get("urls") or not isinstance(data["urls"], list):
            return jsonify({"status": "error", "message": "Missing or invalid 'urls' field. Must be a list of URLs"}), 400

        # Extract and validate data
        knowledge_base_name = data["knowledge_base_name"].strip()
        description = data.get("description", "").strip()
        urls = [url.strip() for url in data["urls"] if url.strip()]
        chunk_size = int(data.get("chunk_size", 1024))
        chunk_overlap = int(data.get("chunk_overlap", 70))
        max_crawl_pages = str(data.get("max_crawl_pages", 10))
        max_crawl_depth = str(data.get("max_crawl_depth", 2))
        dynamic_wait = str(data.get("dynamic_wait", 1))
        provider = data.get("provider", "").strip()
        embedding_model = data.get("embeddingModel", "").strip()
        secret_key = data.get("secretKey", "").strip()

        if not urls:
            return jsonify({"status": "error", "message": "No valid URLs provided"}), 400

        # Validate URLs
        invalid_urls = [url for url in urls if not validate_url(url)]
        if invalid_urls:
            return jsonify({
                "status": "error", 
                "message": f"Invalid URL format(s): {', '.join(invalid_urls)}"
            }), 400

        # Validate chunk parameters
        if chunk_size < 100 or chunk_size > 10000:
            return jsonify({"status": "error", "message": "Chunk size must be between 100 and 10000"}), 400
        if chunk_overlap < 0:
            return jsonify({"status": "error", "message": "Chunk overlap must be non-negative"}), 400

        # Get embedding configuration
        try:
            embedding_config = get_embedding_model(provider, embedding_model, secret_key)
        except ValueError as e:
            logger.error(f"Embedding model initialization failed: {str(e)}")
            return jsonify({"status": "error", "message": str(e)}), 400

        # Process URLs
        logger.info(f"Starting bulk URL processing for {len(urls)} URLs")
        url_chunks, processed_urls, failed_urls = process_multiple_urls(
            urls, knowledge_base_name, chunk_size, chunk_overlap,
            max_crawl_pages, max_crawl_depth, dynamic_wait
        )

        if not url_chunks:
            return jsonify({
                "status": "error", 
                "message": "No content could be extracted from any of the provided URLs",
                "failed_urls": failed_urls
            }), 400

        # Create collection and insert chunks
        collection_name = f"tenant_{tenant_id}_{knowledge_base_name.lower().replace(' ', '_')}"
        try:
            create_qdrant_collection(collection_name, embedding_config["vector_size"])
            insert_chunks_to_qdrant(collection_name, url_chunks, embedding_config)
        except ValueError as e:
            logger.error(f"Qdrant dimension error: {str(e)}")
            return jsonify({"status": "error", "message": str(e)}), 400
        except RuntimeError as e:
            logger.error(f"Qdrant operation failed: {str(e)}")
            return jsonify({"status": "error", "message": str(e)}), 500

        # Save to database
        session = next(db_session())
        try:
            new_knowledge_base = KnowledgeBase(
                tenant_id=tenant_id,
                knowledge_base_name=knowledge_base_name,
                description=description or None,
                upload_pdf=None,
                scrap_url=", ".join(processed_urls),
                max_crawl_pages=int(max_crawl_pages),
                max_crawl_depth=int(max_crawl_depth),
                dynamic_wait=int(dynamic_wait),
                raw_text=None,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                collection_name=collection_name
            )
            session.add(new_knowledge_base)
            session.commit()
            
            response_data = {
                "status": "success",
                "message": "Bulk URL knowledge base created successfully",
                "data": {
                    "knowledge_base_id": new_knowledge_base.knowledge_base_id,
                    "collection_name": collection_name,
                    "processed_urls": processed_urls,
                    "failed_urls": failed_urls,
                    "total_chunks": len(url_chunks),
                    "processing_summary": {
                        "total_urls_requested": len(urls),
                        "successfully_processed": len(processed_urls),
                        "failed": len(failed_urls),
                        "success_rate": f"{(len(processed_urls)/len(urls)*100):.1f}%"
                    }
                }
            }
            
            if failed_urls:
                response_data["warnings"] = f"{len(failed_urls)} URLs failed to process"
            
            logger.info(f"Bulk URL Knowledge Base created: ID={new_knowledge_base.knowledge_base_id}, processed_urls={len(processed_urls)}, chunks={len(url_chunks)}")
            return jsonify(response_data), 201
        except Exception as e:
            session.rollback()
            logger.error(f"Error during bulk URL Knowledge Base creation: {str(e)}")
            return jsonify({"status": "error", "message": "Failed to save Knowledge Base to database"}), 500
        finally:
            session.close()
    except Exception as e:
        logger.error(f"Unexpected error in bulk_url_upload: {str(e)}")
        return jsonify({"status": "error", "message": f"Internal server error: {str(e)}"}), 500

@knowledge_base_blueprint.route("/test-url", methods=["POST"])
@jwt_required()
def test_url_access():
    """Test endpoint to check if a URL can be accessed and processed."""
    try:
        data = request.get_json()
        if not data or not data.get("url"):
            return jsonify({"status": "error", "message": "URL is required"}), 400
        
        url = data["url"].strip()
        if not validate_url(url):
            return jsonify({"status": "error", "message": "Invalid URL format"}), 400
        
        # Test basic URL accessibility
        try:
            response = requests.get(url, timeout=10)
            basic_access = {
                "accessible": True,
                "status_code": response.status_code,
                "content_length": len(response.text),
                "content_type": response.headers.get("content-type", "unknown")
            }
        except Exception as e:
            basic_access = {
                "accessible": False,
                "error": str(e)
            }
        
        # Test Firecrawl processing
        success, result, crawl_stats = crawl_webpage(url, "5", "1", "1")
        firecrawl_test = {
            "success": success,
            "extracted_length": len(result) if success else 0,
            "error": result if not success else None,
            "stats": crawl_stats
        }
        
        return jsonify({
            "status": "success",
            "url": url,
            "basic_access": basic_access,
            "firecrawl_processing": firecrawl_test
        }), 200
        
    except Exception as e:
        logger.error(f"Error in test_url_access: {str(e)}")
        return jsonify({"status": "error", "message": f"Test failed: {str(e)}"}), 500

@knowledge_base_blueprint.route("/debug-form", methods=["POST"])
@jwt_required()
def debug_form_data():
    """Debug endpoint to see exactly what form data is being received."""
    try:
        logger.info("=== DEBUG FORM DATA ENDPOINT ===")
        logger.info(f"Content-Type: {request.content_type}")
        logger.info(f"Method: {request.method}")
        
        # Log all form data
        form_data = {}
        for key, value in request.form.items():
            form_data[key] = {
                "value": value,
                "length": len(value),
                "type": type(value).__name__
            }
            logger.info(f"Form field '{key}': '{value}' (length: {len(value)})")
        
        # Log all files
        files_data = {}
        for key, file_list in request.files.lists():
            files_data[key] = []
            for file in file_list:
                if file.filename:
                    files_data[key].append({
                        "filename": file.filename,
                        "content_type": file.content_type
                    })
        
        return jsonify({
            "status": "success",
            "message": "Form data debugging complete",
            "data": {
                "form_fields": form_data,
                "files": files_data,
                "content_type": request.content_type,
                "method": request.method
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error in debug_form_data: {str(e)}")
        return jsonify({"status": "error", "message": f"Debug failed: {str(e)}"}), 500
        
@knowledge_base_blueprint.route("/<int:kb_id>", methods=["GET"])
@jwt_required()
def list_knowledge_bases_by_kb_id(kb_id):
    """Retrieve a specific knowledge base for the authenticated tenant by ID."""
    try:
        identity = get_jwt_identity()
        claims = get_jwt()
        tenant_id = claims.get("tenant_id")

        if not tenant_id:
            logger.error("Tenant ID not found in token")
            return jsonify({"status": "error", "message": "Tenant ID not found in token"}), 401

        session = next(db_session())
        try:
            # kb_data = (
            #     session.query(CustomBot)
            #     .filter_by(
            #         tenant_id=tenant_id,
            #         bot_id=botId,
            #         del_flg=False
            #     )
            #     .first()
            # )
            
            # kb_id = getattr(kb_data, "knowledge_base_id", None)


            # Fetch knowledge base by tenant and ID
            kb = (
                session.query(KnowledgeBase)
                .filter_by(knowledge_base_id=kb_id, tenant_id=tenant_id, del_flg=False)
                .first()
            )


            if kb :
                knowledge_base_data = {
                    "knowledge_base_id": kb.knowledge_base_id,
                    "tenant_id": kb.tenant_id,
                    "knowledge_base_name": kb.knowledge_base_name,
                    "description": kb.description,
                    "knowledge_base_summary": kb.kb_summary,
                    "upload_pdf": kb.upload_pdf,
                    "scrap_url": kb.scrap_url,
                    "max_crawl_pages": kb.max_crawl_pages,
                    "max_crawl_depth": kb.max_crawl_depth,
                    "dynamic_wait": kb.dynamic_wait,
                    "raw_text": kb.raw_text,
                    "chunk_size": kb.chunk_size,
                    "chunk_overlap": kb.chunk_overlap,
                    "collection_name": kb.collection_name,
                    "created_at": kb.created_at.isoformat() if kb.created_at else None,
                    "updated_at": kb.updated_at.isoformat() if kb.updated_at else None
                }
    
                return jsonify({
                    "data": knowledge_base_data,
                    "status": "success",
                    "message": "Knowledge base retrieved successfully"
                }), 200
            
            if not kb:
                return jsonify({"status": "error", "message": "Knowledge base not found"}), 404

        finally:
            session.close()

    except Exception as e:
        logger.error(f"Unexpected error in list_knowledge_bases: {str(e)}")
        return jsonify({
            "data": [],
            "status": "error",
            "message": f"Internal server error: {str(e)}"
        }), 500

@knowledge_base_blueprint.route("/kb/<int:botId>", methods=["GET"])
@jwt_required()
def list_knowledge_bases_by_bot_id(botId):
    """Retrieve a specific knowledge base for the authenticated tenant by ID."""
    try:
        identity = get_jwt_identity()
        claims = get_jwt()
        tenant_id = claims.get("tenant_id")

        if not tenant_id:
            logger.error("Tenant ID not found in token")
            return jsonify({"status": "error", "message": "Tenant ID not found in token"}), 401

        session = next(db_session())
        try:
            kb_data = (
                session.query(CustomBot)
                .filter_by(
                    tenant_id=tenant_id,
                    bot_id=botId,
                    del_flg=False
                )
                .first()
            )
            
            kb_id = getattr(kb_data, "knowledge_base_id", None)


            # Fetch knowledge base by tenant and ID
            kb = (
                session.query(KnowledgeBase)
                .filter_by(knowledge_base_id=kb_id, tenant_id=tenant_id, del_flg=False)
                .first()
            )


            if kb :
                knowledge_base_data = {
                    "knowledge_base_id": kb.knowledge_base_id,
                    "tenant_id": kb.tenant_id,
                    "knowledge_base_name": kb.knowledge_base_name,
                    "description": kb.description,
                    "upload_pdf": kb.upload_pdf,
                    "scrap_url": kb.scrap_url,
                    "max_crawl_pages": kb.max_crawl_pages,
                    "max_crawl_depth": kb.max_crawl_depth,
                    "dynamic_wait": kb.dynamic_wait,
                    "raw_text": kb.raw_text,
                    "chunk_size": kb.chunk_size,
                    "chunk_overlap": kb.chunk_overlap,
                    "collection_name": kb.collection_name,
                    "created_at": kb.created_at.isoformat() if kb.created_at else None,
                    "updated_at": kb.updated_at.isoformat() if kb.updated_at else None
                }
    
                return jsonify({
                    "data": knowledge_base_data,
                    "status": "success",
                    "message": "Knowledge base retrieved successfully"
                }), 200
            
            if not kb:
                return jsonify({"status": "error", "message": "Knowledge base not found"}), 404

        finally:
            session.close()

    except Exception as e:
        logger.error(f"Unexpected error in list_knowledge_bases: {str(e)}")
        return jsonify({
            "data": [],
            "status": "error",
            "message": f"Internal server error: {str(e)}"
        }), 500
        
# @knowledge_base_blueprint.route("/update/<int:kb_id>", methods=["POST"])
# @jwt_required()
# def update_knowledge_base(kb_id):
#     """Update a new knowledge base with multiple files, web content, and raw text."""
#     try:
#         # Log incoming request details
#         logger.debug(f"Request headers: {request.headers}")
#         logger.debug(f"Content-Type: {request.content_type}")

#         # Validate JWT and tenant ID
#         identity = get_jwt_identity()
#         claims = get_jwt()
#         tenant_id = claims.get("tenant_id")
#         if not tenant_id:
#             logger.error("Tenant ID not found in token")
#             return jsonify({"status": "error", "message": "Tenant ID not found in token"}), 401

#         # Validate Content-Type
#         if not request.content_type or 'multipart/form-data' not in request.content_type:
#             logger.error("Invalid Content-Type: Expected multipart/form-data")
#             return jsonify({"status": "error", "message": "Content-Type must be multipart/form-data"}), 400

#         # Access form data and fi   les
#         try:
#             data = request.form
#             files = request.files.getlist("pdf_files")
#         except Exception as e:
#             logger.error(f"Error accessing form data or files: {str(e)}")
#             return jsonify({"status": "error", "message": "Invalid or corrupted form data"}), 400

#         # Validate required fields
#         required_fields = ["knowledge_base_name"]
#         if not all(field in data and data[field].strip() for field in required_fields):
#             logger.warning("Missing required field: knowledge_base_name")
#             return jsonify({"status": "error", "message": "Missing required field: knowledge_base_name"}), 400

#         # Extract form data
#         knowledge_base_name = data["knowledge_base_name"].strip()
#         scrap_urls = [data[key].strip() for key in data if key.startswith("scrap_urls") and data[key].strip()]
#         raw_text = data.get("raw_text", "").strip()
#         provider = data.get("provider", "").strip()
#         embedding_model = data.get("embeddingModel", "").strip()
#         secret_key = data.get("secretKey", "").strip()
#         chunk_size = int(data.get("chunk_size", 1024))
#         chunk_overlap = int(data.get("chunk_overlap", 50))

#         # Validate chunk size and overlap
#         if chunk_size < 100 or chunk_size > 10000:
#             logger.warning(f"Invalid chunk_size: {chunk_size}")
#             return jsonify({"status": "error", "message": "Chunk size must be between 100 and 10000"}), 400
#         if chunk_overlap < 0:
#             logger.warning(f"Invalid chunk_overlap: {chunk_overlap}")
#             return jsonify({"status": "error", "message": "Chunk overlap must be non-negative"}), 400

#         # Validate file types
#         allowed_extensions = {'.pdf', '.docx', '.txt'}
#         invalid_files = [file.filename for file in files if os.path.splitext(file.filename)[1].lower() not in allowed_extensions]
#         if invalid_files:
#             logger.warning(f"Invalid file types uploaded: {invalid_files}")
#             return jsonify({"status": "error", "message": f"Only PDF, DOCX, and TXT files are allowed: {invalid_files}"}), 400

#         # Get embedding configuration
#         try:
#             embedding_config = get_embedding_model(provider, embedding_model, secret_key)
#         except ValueError as e:
#             logger.error(f"Embedding model initialization failed: {str(e)}")
#             return jsonify({"status": "error", "message": str(e)}), 400

#         # Process files and collect chunks
#         chunks = []
#         processed_files = []
#         upload_dir = "uploads"
#         os.makedirs(upload_dir, exist_ok=True)

#         for file in files:
#             if not file.filename:
#                 logger.warning("Empty filename in uploaded files")
#                 continue
#             file_ext = os.path.splitext(file.filename)[1].lower()
#             pdf_name = os.path.splitext(file.filename)[0]
#             pdf_path = os.path.join(upload_dir, file.filename)

#             try:
#                 file.save(pdf_path)
#                 logger.info(f"Saved file: {file.filename}")
#             except Exception as e:
#                 logger.error(f"Failed to save file {file.filename}: {str(e)}")
#                 return jsonify({"status": "error", "message": f"Failed to save file {file.filename}: {str(e)}"}), 500

#             output_dir = create_directory_structure(knowledge_base_name, pdf_name)
#             if file_ext == '.pdf':
#                 success, message, file_chunks = process_pdf(pdf_path, output_dir, knowledge_base_name, chunk_size, chunk_overlap)
#                 if not success:
#                     logger.error(f"Failed to process PDF {file.filename}: {message}")
#                     return jsonify({"status": "error", "message": f"Failed to process PDF {file.filename}: {message}"}), 500
#                 chunks.extend(file_chunks)
#                 processed_files.append(file.filename)
#             else:
#                 # Handle DOCX and TXT files (assuming text extraction logic is needed)
#                 try:
#                     if file_ext == '.docx':
#                         from docx import Document
#                         doc = Document(pdf_path)
#                         text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#                     elif file_ext == '.txt':
#                         with open(pdf_path, 'r', encoding='utf-8') as f:
#                             text = f.read()
#                     file_chunks = chunk_text(text, chunk_size, chunk_overlap)
#                     chunks.extend(file_chunks)
#                     processed_files.append(file.filename)
#                     logger.info(f"Processed {file_ext} file: {file.filename}")
#                 except Exception as e:
#                     logger.error(f"Failed to process {file_ext} file {file.filename}: {str(e)}")
#                     return jsonify({"status": "error", "message": f"Failed to process {file_ext} file {file.filename}: {str(e)}"}), 500

#         # Process web-scraped content
#         # if scrap_url:
#         #     success, extracted_text = crawl_webpage(scrap_url, data.get("max_crawl_pages"), data.get("max_crawl_depth"), data.get("dynamic_wait"))
#         #     if not success:
#         #         logger.error(f"Web crawling failed: {extracted_text}")
#         #         return jsonify({"status": "error", "message": extracted_text}), 500
#         #     website_chunks = chunk_text(extracted_text, chunk_size, chunk_overlap)
#         #     chunks.extend(website_chunks)
#         #     create_website_chunks_directory(knowledge_base_name, scrap_url)
#         #     logger.info(f"Processed web content from {scrap_url}")

#         if scrap_urls:
#             with ThreadPoolExecutor(max_workers=min(len(scrap_urls), 5)) as executor:
#                 future_to_url = {
#                     executor.submit(
#                         crawl_webpage,
#                         url,
#                         data.get("max_crawl_pages"),
#                         data.get("max_crawl_depth"),
#                         data.get("dynamic_wait")
#                     ): url
#                     for url in scrap_urls
#                 }

#                 for future in as_completed(future_to_url):
#                     url = future_to_url[future]
#                     try:
#                         success, extracted_text = future.result()
#                         scrap_url = url  # reuse the known URL
#                         if not success:
#                             logger.error(f"Web crawling failed for {url}: {extracted_text}")
#                             return jsonify({"status": "error", "message": f"Web crawling failed for {url}: {extracted_text}"}), 500
#                         website_chunks = chunk_text(extracted_text, chunk_size, chunk_overlap)
#                         chunks.extend(website_chunks)
#                         create_website_chunks_directory(knowledge_base_name, url)
#                         logger.info(f"Processed web content from {url}")
    
#                     except Exception as e:
#                         logger.error(f"Exception processing {url}: {str(e)}")
#                         return jsonify({"status": "error", "message": f"Exception processing {url}: {str(e)}"}), 500

#         # Process raw text
#         if raw_text:
#             text_chunks = chunk_text(raw_text, chunk_size, chunk_overlap)
#             chunks.extend(text_chunks)
#             create_raw_text_chunks_directory(knowledge_base_name)
#             logger.info("Processed raw text")

#         # Validate chunks
#         if not chunks:
#             logger.warning("No chunks generated from files, web, or raw text")
#             return jsonify({"status": "error", "message": "No content provided to create knowledge base"}), 400

#         # Create collection and insert chunks
#         collection_name = f"tenant_{tenant_id}_{knowledge_base_name.lower().replace(' ', '_')}"
#         try:
#             create_qdrant_collection(collection_name, embedding_config["vector_size"])
#             insert_chunks_to_qdrant(collection_name, chunks, embedding_config)
#         except ValueError as e:
#             logger.error(f"Qdrant dimension error: {str(e)}")
#             return jsonify({"status": "error", "message": str(e)}), 400
#         except RuntimeError as e:
#             logger.error(f"Qdrant operation failed: {str(e)}")
#             return jsonify({"status": "error", "message": str(e)}), 500

#         # Save to database
#         session = next(db_session())
#         try:
#             # Find the existing record for this tenant and kb_id
#             existing_kb = session.query(KnowledgeBase).filter_by(
#                 tenant_id=tenant_id,
#                 knowledge_base_id=kb_id
#             ).first()

#             if not existing_kb:
#                 logger.warning(f"KnowledgeBase with ID {kb_id} not found for tenant {tenant_id}")
#                 return jsonify({"status": "error", "message": "Knowledge base not found"}), 404

#             # Update fields
#             existing_kb.knowledge_base_name = knowledge_base_name
#             existing_kb.upload_pdf = ", ".join(processed_files) if processed_files else existing_kb.upload_pdf
#             existing_kb.scrap_url = scrap_url if scrap_url else existing_kb.scrap_url
#             existing_kb.max_crawl_pages = int(data.get("max_crawl_pages", 0)) or None
#             existing_kb.max_crawl_depth = int(data.get("max_crawl_depth", 0)) or None
#             existing_kb.dynamic_wait = int(data.get("dynamic_wait", 0)) or None
#             existing_kb.raw_text = raw_text if raw_text else existing_kb.raw_text
#             existing_kb.chunk_size = chunk_size
#             existing_kb.chunk_overlap = chunk_overlap
#             existing_kb.collection_name = collection_name

#             # Commit update
#             session.commit()

#             response_data = {
#                 "status": "success",
#                 "message": "KnowledgeBase updated successfully",
#                 "data": {
#                     "knowledge_base_id": existing_kb.knowledge_base_id,
#                     "collection_name": collection_name
#                 }
#             }
#             logger.info(f"KnowledgeBase updated: knowledge_base_id={existing_kb.knowledge_base_id}, collection_name={collection_name}")
#             return jsonify(response_data), 200

#         except Exception as e:
#             session.rollback()
#             logger.error(f"Error during KnowledgeBase creation: {str(e)}")
#             return jsonify({"status": "error", "message": "Failed to save KnowledgeBase to database"}), 500
#         finally:
#             session.close()
#     except Exception as e:
#         logger.error(f"Unexpected error in create_knowledge_base: {str(e)}")
#         return jsonify({"status": "error", "message": f"Internal server error: {str(e)}"}), 500


@knowledge_base_blueprint.route("/update/<int:kb_id>", methods=["POST"])
@jwt_required()
def update_knowledge_base(kb_id):
    """Update an existing knowledge base with new files, URLs, or raw text."""
    try:
        logger.info("=== KNOWLEDGE BASE UPDATE REQUEST ===")
        logger.debug(f"Content-Type: {request.content_type}")

        # Validate JWT
        claims = get_jwt()
        tenant_id = claims.get("tenant_id")
        if not tenant_id:
            return jsonify({"status": "error", "message": "Tenant ID not found in token"}), 401

        # Check form-data content type
        if "multipart/form-data" not in request.content_type:
            return jsonify({"status": "error", "message": "Content-Type must be multipart/form-data"}), 400

        # Read form data
        data = request.form

        # Receive all file types (FIXED)
        pdf_files = request.files.getlist("pdf_files")
        excel_files = request.files.getlist("excel_files")
        docx_files = request.files.getlist("docx_files")
        txt_files = request.files.getlist("txt_files")

        files = pdf_files + excel_files + docx_files + txt_files

        logger.info(f"Files received: {[f.filename for f in files]}")

        # Validate required field
        if not data.get("knowledge_base_name"):
            return jsonify({"status": "error", "message": "Missing required field: knowledge_base_name"}), 400

        # Extract form fields
        knowledge_base_name = data["knowledge_base_name"].strip()
        description = data.get("description", "").strip()
        scrap_urls = request.form.getlist("scrap_urls")
        scrap_urls = [url.strip() for url in scrap_urls if url.strip()]
        raw_text = data.get("raw_text", "").strip()
        chunk_size = int(data.get("chunk_size", 1024))
        chunk_overlap = int(data.get("chunk_overlap", 70))

        provider = data.get("provider", "").strip()
        embedding_model = data.get("embeddingModel", "").strip()
        secret_key = data.get("secretKey", "").strip()

        # Validate embedding config
        try:
            embedding_config = get_embedding_model(provider, embedding_model, secret_key)
        except ValueError as e:
            return jsonify({"status": "error", "message": str(e)}), 400

        # Validate file extensions (FIXED)
        allowed_extensions = {'.pdf', '.docx', '.txt', '.xls', '.xlsx'}
        invalid_files = [
            f.filename for f in files
            if f.filename and os.path.splitext(f.filename)[1].lower() not in allowed_extensions
        ]
        if invalid_files:
            return jsonify({
                "status": "error",
                "message": f"Invalid file types: {invalid_files}. Allowed: PDF, DOCX, TXT, XLS, XLSX"
            }), 400

        # Variables for processed content
        chunks = []
        processed_files = []
        upload_dir = "Uploads"
        os.makedirs(upload_dir, exist_ok=True)

        # Process uploaded files (ALL types supported)
        for file in files:
            if not file.filename:
                continue

            filename = file.filename
            file_ext = os.path.splitext(filename)[1].lower()
            file_path = os.path.join(upload_dir, filename)

            file.save(file_path)
            logger.info(f"Saved: {filename}")

            # --- EXCEL FILES ---
            if file_ext in ['.xls', '.xlsx']:
                logger.info(f"[UPDATE] Processing Excel: {filename}")

                success, message, excel_chunks = process_excel_file(
                    file_path=file_path,
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap
                )

                if not success:
                    return jsonify({"status": "error", "message": message}), 400

                chunks.extend(excel_chunks)
                processed_files.append(filename)
                continue

            # --- PDF FILES ---
            if file_ext == ".pdf":
                output_dir = create_directory_structure(knowledge_base_name, filename.split(".")[0])
                success, message, pdf_chunks = process_pdf(
                    file_path, output_dir, knowledge_base_name, chunk_size, chunk_overlap
                )

                if not success:
                    return jsonify({"status": "error", "message": message}), 400

                chunks.extend([{"text": c, "source_type": "pdf"} for c in pdf_chunks])
                processed_files.append(filename)
                continue

            # --- DOCX FILES ---
            if file_ext == ".docx":
                text = process_docx_file(file_path)
                docx_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type="docx")
                chunks.extend(docx_chunks)
                processed_files.append(filename)
                continue

            # --- TXT FILES ---
            if file_ext == ".txt":
                text = process_txt_file(file_path)
                txt_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type="txt")
                chunks.extend(txt_chunks)
                processed_files.append(filename)
                continue

        # --- PROCESS SCRAP URLS ---
        processed_urls = []
        failed_urls = []

        if scrap_urls:
            for url in scrap_urls:
                logger.info(f"Processing URL: {url}")

                # Skip LinkedIn
                if "linkedin.com" in url:
                    failed_urls.append({"url": url, "error": "LinkedIn scraping not supported"})
                    continue

                success, extracted, _ = crawl_webpage(url, "", "", "")

                if success and extracted.strip():
                    logger.info(f"SUCCESS: {url}")

                    url_chunks = chunk_text(
                        extracted,
                        chunk_size,
                        chunk_overlap,
                        source_type="web"
                    )

                    chunks.extend(url_chunks)
                    processed_urls.append(url)
                else:
                    logger.error(f"FAILED: {url} | Reason: {extracted}")
                    failed_urls.append({"url": url, "error": extracted})

        # --- RAW TEXT ---
        if raw_text:
            raw_chunks = chunk_text(raw_text, chunk_size, chunk_overlap, source_type="raw")
            chunks.extend(raw_chunks)

        # Ensure content exists
        if not chunks:
            return jsonify({
                "status": "error",
                "message": "All URLs failed or no content extracted",
                "failed_urls": failed_urls
            }), 400
        # Create Qdrant collection & insert chunks
        collection_name = f"tenant_{tenant_id}_{knowledge_base_name.lower().replace(' ', '_')}"

        create_qdrant_collection(collection_name, embedding_config["vector_size"])
        insert_chunks_to_qdrant(collection_name, chunks, embedding_config)

        # Update DB record
        session = next(db_session())

        kb = session.query(KnowledgeBase).filter_by(
            tenant_id=tenant_id,
            knowledge_base_id=kb_id
        ).first()

        if not kb:
            return jsonify({"status": "error", "message": "Knowledge base not found"}), 404

        kb.knowledge_base_name = knowledge_base_name
        kb.description = description if description else kb.description
        kb.upload_pdf = ", ".join(processed_files) if processed_files else kb.upload_pdf
        kb.scrap_url = ", ".join(processed_urls) if processed_urls else kb.scrap_url
        kb.raw_text = raw_text if raw_text else kb.raw_text
        kb.chunk_size = chunk_size
        kb.chunk_overlap = chunk_overlap
        kb.collection_name = collection_name

        session.commit()

        return jsonify({
            "status": "success",
            "message": "KnowledgeBase updated successfully",
            "data": {
                "knowledge_base_id": kb.knowledge_base_id,
                "collection_name": collection_name,
                "processed_files": processed_files,
                "processed_urls": processed_urls,
                "failed_urls": failed_urls,
                "total_chunks": len(chunks)
            }
        }), 200

    except Exception as e:
        logger.error(f"Unexpected error in update_knowledge_base: {e}")
        return jsonify({"status": "error", "message": f"Internal server error: {str(e)}"}), 500


@knowledge_base_blueprint.route("/delete/<int:kb_id>", methods=["DELETE"])
@jwt_required()
def delete_knowledge_base(kb_id):
    """Delete a knowledge base by ID for the authenticated tenant."""
    try:
        # Get tenant_id from JWT
        claims = get_jwt()
        tenant_id = claims.get("tenant_id")
        if not tenant_id:
            logger.error("Tenant ID not found in token")
            return jsonify({"status": "error", False: "Tenant ID not found in token"}), 401

        # Open DB session
        session = next(db_session())

        # Find the knowledge base record
        kb_record = session.query(KnowledgeBase).filter_by(
            tenant_id=tenant_id,
            knowledge_base_id=kb_id
        ).first()

        if not kb_record:
            logger.warning(f"KnowledgeBase with ID {kb_id} not found for tenant {tenant_id}")
            return jsonify({"status": False, "message": "Knowledge base not found"}), 404

        # Delete the record
        session.delete(kb_record)
        session.commit()

        logger.info(f"KnowledgeBase deleted: ID={kb_id}, tenant={tenant_id}")
        return jsonify({"status":True, "message": "Knowledge base deleted successfully"}), 200

    except Exception as e:
        logger.error(f"Error deleting knowledge base {kb_id}: {str(e)}")
        return jsonify({"status":False, "message": f"Internal server error: {str(e)}"}), 500
    finally:
        try:
            session.close()
        except:
            pass


def process_knowledge_base_async(app, task_id: str, kb_id: int, tenant_id: int, metadata: dict):
    """
    Background worker function to process KB creation.
    This runs in a separate thread and updates the build task status as it progresses.
    """
    from app.services.build_tasks_manager import get_build_tasks_manager
    from datetime import datetime
    import shutil
    
    manager = get_build_tasks_manager()
    
    with app.app_context():
        try:
            # Start the task
            manager.start_task(task_id)
            manager.update_progress(task_id, 5, "Initializing...", "Starting KB build process")
            
            session = next(db_session())
            
            try:
                # Update KB status in database
                kb = session.query(KnowledgeBase).filter_by(knowledge_base_id=kb_id).first()
                if not kb:
                    raise Exception(f"Knowledge base {kb_id} not found")
                
                logger.info("[KB BUILD] start | task_id=%s kb_id=%s tenant_id=%s", task_id, kb_id, tenant_id)
                kb.status = "processing"
                kb.started_at = datetime.utcnow()
                kb.error_message = None
                session.commit()
                
                # Extract parameters from metadata
                upload_dir = metadata.get("upload_dir")
                files_info = metadata.get("files", [])
                scrap_urls = metadata.get("scrap_urls", [])
                raw_text = metadata.get("raw_text", "")
                provider = metadata.get("provider", "")
                embedding_model = metadata.get("embedding_model", "")
                secret_key = metadata.get("secret_key", "")
                chunk_size = metadata.get("chunk_size", 1024)
                chunk_overlap = metadata.get("chunk_overlap", 150)
                max_crawl_pages = metadata.get("max_crawl_pages", "")
                max_crawl_depth = metadata.get("max_crawl_depth", "")
                dynamic_wait = metadata.get("dynamic_wait", "")
                removed_files = metadata.get("removed_files", [])
                
                knowledge_base_name = kb.knowledge_base_name
                
                manager.update_progress(task_id, 10, "Validating embedding model...", "Checking embedding configuration")
                
                # Validate embedding model
                try:
                    embedding_config = get_embedding_model(provider, embedding_model, secret_key)
                except ValueError as e:
                    raise Exception(f"Embedding model error: {str(e)}")
                
                chunks = []
                processed_files = []
                processed_urls = []
                failed_urls = []
                media_type = None

                # Process files
                total_files = len(files_info)
                processed_content_info = []  # Track content details for View Content feature
                
                if total_files > 0:
                    manager.update_progress(task_id, 15, f"Processing {total_files} document(s)...", f"📄 Starting document processing: {total_files} file(s)")
                    
                    for i, file_info in enumerate(files_info):
                        # Check if task was cancelled
                        task = manager.get_task(task_id)
                        if task and task.get("status") == "cancelled":
                            raise Exception("Build cancelled by user")
                        
                        file_path = file_info.get("path")
                        filename = file_info.get("filename")
                        file_ext = os.path.splitext(filename)[1].lower()
                        file_name_without_ext = os.path.splitext(filename)[0]
                        
                        progress = 15 + int(((i + 1) / total_files) * 35)
                        manager.update_progress(task_id, progress, f"Processing document {i+1}/{total_files}: {filename}", f"📄 Document {i+1}/{total_files}: {filename}")
                        
                        chunks_before = len(chunks)
                        
                        try:
                            if file_ext in ['.xls', '.xlsx']:
                                success, message, excel_chunks = process_excel_file(
                                    file_path=file_path,
                                    chunk_size=chunk_size,
                                    chunk_overlap=chunk_overlap
                                )
                                if not success:
                                    logger.error(f"Excel processing failed: {message}")
                                    failed_urls.append({"file": filename, "error": message})
                                    manager.update_progress(task_id, progress, f"Processing document {i+1}/{total_files}", f"❌ Failed: {filename} - {message}")
                                    continue
                                chunks.extend(excel_chunks)
                                processed_files.append(filename)
                                media_type = "excel"
                                
                            elif file_ext == '.pdf':
                                output_dir = create_directory_structure(knowledge_base_name, file_name_without_ext)
                                success, message, file_chunks = process_pdf(file_path, output_dir, knowledge_base_name, chunk_size, chunk_overlap)
                                if not success:
                                    logger.error(f"PDF processing failed: {message}")
                                    failed_urls.append({"file": filename, "error": message})
                                    manager.update_progress(task_id, progress, f"Processing document {i+1}/{total_files}", f"❌ Failed: {filename} - {message}")
                                    continue
                                chunks.extend([{"text": chunk, "source_type": "pdf", "source_file": filename} for chunk in file_chunks])
                                processed_files.append(filename)

                            elif file_ext == '.docx' and DOCX_AVAILABLE:
                                text = process_docx_file(file_path)
                                file_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type='docx')
                                chunks.extend([{**c, "source_file": filename} for c in file_chunks])
                                processed_files.append(filename)

                            elif file_ext == '.txt':
                                text = process_txt_file(file_path)
                                file_chunks = chunk_text(text, chunk_size, chunk_overlap, source_type='txt')
                                chunks.extend([{**c, "source_file": filename} for c in file_chunks])
                                processed_files.append(filename)
                            
                            # Log success with chunk count
                            chunks_created = len(chunks) - chunks_before
                            manager.update_progress(task_id, progress, f"Processing document {i+1}/{total_files}", f"✅ Processed: {filename} → {chunks_created} chunks")
                            
                            # Store content info for View Content feature
                            processed_content_info.append({
                                "type": "document",
                                "name": filename,
                                "file_type": file_ext[1:].upper(),
                                "chunks_created": chunks_created,
                                "sample": chunks[-1].get("text", "")[:200] if chunks and isinstance(chunks[-1], dict) else ""
                            })
                                
                        except Exception as e:
                            logger.error(f"Error processing {filename}: {str(e)}")
                            failed_urls.append({"file": filename, "error": str(e)})
                            manager.update_progress(task_id, progress, f"Processing document {i+1}/{total_files}", f"❌ Error: {filename} - {str(e)}")
                
                # Process URLs
                if scrap_urls:
                    total_urls = len(scrap_urls)
                    manager.update_progress(task_id, 55, f"Crawling {total_urls} URL(s)...", f"🌐 Starting web crawl: {total_urls} URL(s)")
                    
                    from concurrent.futures import ThreadPoolExecutor, as_completed
                    
                    urls_processed = 0
                    total_pages_discovered = 0
                    total_pages_crawled = 0
                    total_credits_used = 0
                    crawl_skips = []
                    redirect_chains = []
                    selected_modes = {}
                    total_duplicates_removed = 0
                    
                    with ThreadPoolExecutor(max_workers=min(len(scrap_urls), 5)) as executor:
                        future_to_url = {
                            executor.submit(crawl_webpage, url, max_crawl_pages, max_crawl_depth, dynamic_wait): url
                            for url in scrap_urls
                        }
                        for future in as_completed(future_to_url):
                            url = future_to_url[future]
                            urls_processed += 1
                            chunks_before = len(chunks)
                            
                            try:
                                success, extracted_text, crawl_stats = future.result()
                                pages_discovered = crawl_stats.get("pages_discovered", 0)
                                pages_crawled = crawl_stats.get("pages_crawled", 0)
                                credits_used = crawl_stats.get("credits_used", 0)
                                crawl_mode = crawl_stats.get("crawl_mode", "unknown")
                                crawled_urls_list = crawl_stats.get("crawled_urls", [])
                                skipped_urls_list = crawl_stats.get("skipped_urls", [])
                                redirect_chain = crawl_stats.get("redirect_chain", [])
                                crawl_mode_selected = crawl_stats.get("crawl_mode_selected", "STANDARD")
                                crawl_mode_flags = crawl_stats.get("crawl_mode_flags", [])
                                duplicates_removed = int(crawl_stats.get("duplicates_removed", 0) or 0)
                                
                                total_pages_discovered += pages_discovered
                                total_pages_crawled += pages_crawled
                                total_credits_used += credits_used
                                total_duplicates_removed += duplicates_removed
                                selected_modes[crawl_mode_selected] = selected_modes.get(crawl_mode_selected, 0) + 1
                                if skipped_urls_list:
                                    crawl_skips.extend(skipped_urls_list)
                                if redirect_chain:
                                    redirect_chains.append(redirect_chain)
                                
                                if success:
                                    website_chunks = chunk_web_content_by_source(
                                        extracted_text,
                                        chunk_size,
                                        chunk_overlap,
                                        crawl_meta=crawl_stats,
                                    )
                                    chunks.extend(website_chunks)
                                    processed_urls.append(url)
                                    create_website_chunks_directory(knowledge_base_name, url)
                                    
                                    chunks_created = len(chunks) - chunks_before
                                    char_count = len(extracted_text)
                                    progress = 55 + int((urls_processed / total_urls) * 10)
                                    
                                    # Log with Firecrawl stats
                                    stats_info = f"📊 {pages_crawled} pages, {credits_used} credits" if crawl_mode == "multi" else "📄 single page"
                                    mode_info = f"mode={crawl_mode_selected} flags={','.join(crawl_mode_flags)}"
                                    manager.update_progress(task_id, progress, f"Crawling URL {urls_processed}/{total_urls}", 
                                        f"✅ Crawled: {url[:40]}... → {chunks_created} chunks ({char_count:,} chars) [{stats_info}; {mode_info}]")
                                    
                                    # Store content info for View Content feature
                                    processed_content_info.append({
                                        "type": "url",
                                        "name": url,
                                        "chunks_created": chunks_created,
                                        "char_count": char_count,
                                        "pages_crawled": pages_crawled,
                                        "credits_used": credits_used,
                                        "crawled_urls": crawled_urls_list[:10],  # Store first 10 URLs
                                        "sample": website_chunks[0].get("text", "")[:200] if website_chunks and isinstance(website_chunks[0], dict) else ""
                                    })
                                else:
                                    failed_urls.append({"url": url, "error": extracted_text})
                                    manager.update_progress(task_id, 55 + int((urls_processed / total_urls) * 10), f"Crawling URL {urls_processed}/{total_urls}", f"❌ Failed: {url[:50]}... - {extracted_text[:50]}")
                            except Exception as e:
                                failed_urls.append({"url": url, "error": str(e)})
                                manager.update_progress(task_id, 55 + int((urls_processed / total_urls) * 10), f"Crawling URL {urls_processed}/{total_urls}", f"❌ Error: {url[:50]}... - {str(e)[:50]}")
                    
                    # Log total crawl stats summary
                    if total_pages_crawled > 0:
                        manager.update_progress(task_id, 65, "Web crawl complete", 
                            f"📊 Total crawl: {total_pages_crawled} pages from {urls_processed} URLs ({total_credits_used} credits used)")
                    logger.info(
                        "[KB CRAWL STATS] task_id=%s discovered=%s crawled=%s urls_input=%s urls_success=%s urls_failed=%s skipped=%s redirects=%s duplicates_removed=%s credits=%s modes=%s",
                        task_id,
                        total_pages_discovered,
                        total_pages_crawled,
                        total_urls,
                        len(processed_urls),
                        len([x for x in failed_urls if x.get('url')]),
                        len(crawl_skips),
                        len(redirect_chains),
                        total_duplicates_removed,
                        total_credits_used,
                        selected_modes,
                    )
                    if crawl_skips:
                        logger.info("[KB CRAWL SKIPS] task_id=%s skipped_urls=%s", task_id, crawl_skips[:25])
                    if redirect_chains:
                        logger.info("[KB REDIRECT CHAINS] task_id=%s chains=%s", task_id, redirect_chains[:25])
                
                # Process raw text
                if raw_text:
                    chunks_before = len(chunks)
                    manager.update_progress(task_id, 65, "Processing raw text...", f"📝 Processing raw text ({len(raw_text):,} chars)")
                    text_chunks = chunk_text(raw_text, chunk_size, chunk_overlap, source_type='raw')
                    chunks.extend(text_chunks)
                    create_raw_text_chunks_directory(knowledge_base_name)
                    
                    chunks_created = len(chunks) - chunks_before
                    manager.update_progress(task_id, 67, "Raw text processed", f"✅ Raw text processed → {chunks_created} chunks")
                    
                    # Store content info for View Content feature
                    processed_content_info.append({
                        "type": "raw_text",
                        "name": "Raw Text Input",
                        "chunks_created": chunks_created,
                        "char_count": len(raw_text),
                        "sample": raw_text[:200]
                    })
                
                # Check if we have content
                if not chunks:
                    raise Exception("No content was processed successfully")

                logger.info(
                    "[KB BUILD] chunks created | task_id=%s kb_id=%s chunks=%s",
                    task_id, kb_id, len(chunks)
                )
                
                manager.update_progress(task_id, 70, "Creating vector collection...", f"Creating Qdrant collection with {len(chunks)} chunks")

                # Create Qdrant collection
                collection_name = f"tenant_{tenant_id}_{knowledge_base_name.lower().replace(' ', '_')}"
                try:
                    create_qdrant_collection(collection_name, embedding_config["vector_size"])
                except Exception as e:
                    raise Exception(f"Failed to create collection: {str(e)}")
                
                manager.update_progress(task_id, 75, "Generating embeddings...", f"Processing {len(chunks)} chunks")
                
                # Insert chunks to Qdrant
                try:
                    insert_chunks_to_qdrant(collection_name, chunks, embedding_config)
                except Exception as e:
                    raise Exception(f"Failed to insert embeddings: {str(e)}")

                logger.info(
                    "[KB BUILD] embeddings+vectors inserted | task_id=%s kb_id=%s vectors=%s collection=%s",
                    task_id, kb_id, len(chunks), collection_name
                )
                logger.info(
                    "[KB DB STATS] task_id=%s kb_id=%s chunks_stored=%s embeddings_success_rate=%.2f%%",
                    task_id, kb_id, len(chunks), 100.0
                )

                # Remove deleted source files only after the new build has succeeded.
                # This avoids permanently dropping old vectors when the replacement crawl fails.
                if removed_files:
                    from qdrant_client.models import Filter, FieldCondition, MatchValue
                    target_collection = kb.collection_name or collection_name
                    manager.update_progress(
                        task_id,
                        88,
                        "Syncing removed files...",
                        f"🗑️ Removing {len(removed_files)} file(s) from vector store"
                    )
                    for fname in removed_files:
                        try:
                            qdrant.delete(
                                collection_name=target_collection,
                                points_selector=Filter(
                                    must=[FieldCondition(key="source_file", match=MatchValue(value=fname))]
                                )
                            )
                            logger.info(f"[KB UPDATE] Deleted vectors for removed file: {fname}")
                        except Exception as e:
                            logger.warning(f"[KB UPDATE] Could not delete vectors for {fname}: {e}")
                
                manager.update_progress(task_id, 90, "Generating KB summary...", "Analyzing content for summary")
                
                # Generate KB summary
                from utils import generate_kb_summary_from_chunks
                kb_summary = ""
                try:
                    kb_summary = generate_kb_summary_from_chunks(
                        kb_name=collection_name,
                        qdrant_client=qdrant,
                        limit=6
                    )
                    if kb_summary and len(kb_summary) > 5000:
                        kb_summary = kb_summary[:5000]
                except Exception as e:
                    logger.warning(f"KB summary generation failed: {e}")
                
                manager.update_progress(task_id, 95, "Saving to database...", "Finalizing knowledge base record")
                
                # Update KB record (preserve existing source metadata on partial/empty updates)
                existing_files = [x.strip() for x in (kb.upload_pdf or "").split(",") if x.strip()]
                merged_files = list(dict.fromkeys(existing_files + processed_files))
                kb.upload_pdf = ", ".join(merged_files) if merged_files else kb.upload_pdf
                kb.upload_media = ", ".join(merged_files) if merged_files else kb.upload_media
                kb.media_type = media_type
                existing_urls = [x.strip() for x in (kb.scrap_url or "").split(",") if x.strip()]
                merged_urls = list(dict.fromkeys(existing_urls + processed_urls))
                kb.scrap_url = ", ".join(merged_urls) if merged_urls else kb.scrap_url
                kb.max_crawl_pages = int(max_crawl_pages or 0) or None
                kb.max_crawl_depth = int(max_crawl_depth or 0) or None
                kb.dynamic_wait = int(dynamic_wait or 0) or None
                kb.raw_text = raw_text if raw_text else kb.raw_text
                kb.collection_name = collection_name
                kb.kb_summary = kb_summary or None
                kb.total_chunks = len(chunks)
                kb.processed_chunks = len(chunks)
                kb.status = "completed"
                kb.completed_at = datetime.utcnow()
                kb.error_message = None
                
                session.commit()
                logger.info(
                    "[KB BUILD] db update success | task_id=%s kb_id=%s status=%s total_chunks=%s",
                    task_id, kb_id, kb.status, kb.total_chunks
                )
                
                # Clean up upload directory
                if upload_dir and os.path.exists(upload_dir):
                    try:
                        shutil.rmtree(upload_dir)
                    except Exception as e:
                        logger.warning(f"Failed to clean up upload dir: {e}")
                
                # Mark task as completed with content info
                build_summary = {
                    "total_chunks": len(chunks),
                    "documents_processed": len(processed_files),
                    "urls_crawled": len(processed_urls),
                    "failed_items": len(failed_urls),
                    "collection_name": collection_name,
                    "content_sources": processed_content_info
                }
                manager.complete_task(
                    task_id,
                    knowledge_base_id=kb_id,
                    success_message=f"✅ KB created: {len(chunks)} chunks from {len(processed_files)} docs, {len(processed_urls)} URLs",
                    build_summary=build_summary
                )
                
                logger.info(f"Async KB build completed: task_id={task_id}, kb_id={kb_id}, chunks={len(chunks)}")
                
            except Exception as e:
                session.rollback()
                raise
            finally:
                session.close()
                
        except Exception as e:
            error_msg = str(e)
            logger.error(f"Async KB build failed: task_id={task_id}, error={error_msg}")
            
            # Update task status
            manager.fail_task(task_id, error_msg)
            
            # Update KB status in database
            try:
                session = next(db_session())
                kb = session.query(KnowledgeBase).filter_by(knowledge_base_id=kb_id).first()
                if kb:
                    kb.status = "failed"
                    kb.status = "failed"
                    kb.error_message = error_msg
                    kb.completed_at = datetime.utcnow()
                    session.commit()
                    logger.info(
                        "[KB BUILD] db update failure | task_id=%s kb_id=%s status=%s error=%s",
                        task_id, kb_id, kb.status, error_msg
                    )
                session.close()
            except Exception as db_error:
                logger.error(f"Failed to update KB status: {db_error}")
            
            # Clean up upload directory on failure
            upload_dir = metadata.get("upload_dir")
            if upload_dir and os.path.exists(upload_dir):
                try:
                    shutil.rmtree(upload_dir)
                except:
                    pass


def rebuild_knowledge_base_async(app, task_id: str, kb_id: int, tenant_id: int):
    """
    Rebuild an existing knowledge base (retry failed build).
    """
    from app.services.build_tasks_manager import get_build_tasks_manager
    
    manager = get_build_tasks_manager()
    task = manager.get_task(task_id)
    
    if task:
        metadata = task.get("metadata", {})
        process_knowledge_base_async(app, task_id, kb_id, tenant_id, metadata)
